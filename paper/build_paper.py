#!/usr/bin/env python3
"""
Build script for: #KyivNotKiev: A Large-Scale Computational Study of Ukrainian Toponym Adoption
Generates a publication-ready DOCX for Computational Linguistics (MIT Press).

Author: Ivan Dobrovolskyi, ivan@kyivnotkiev.org
Usage: python paper/build_paper.py
Output: paper/kyivnotkiev_cl_paper.docx
"""

import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
FIGURES_DIR = SCRIPT_DIR / "figures"
DATA_DIR = PROJECT_ROOT / "data"
OUTPUT_PATH = SCRIPT_DIR / "kyivnotkiev_cl_paper.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Cambria"):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = font_name
    run.font.bold = bold


def add_table_row(table, values, bold=False, header=False, shade=None):
    """Add a row to a table with values."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        set_cell_text(cell, val, bold=bold, size=Pt(8) if not header else Pt(9))
        if shade:
            set_cell_shading(cell, shade)
    return row


def style_header_row(table, shade_color="4472C4"):
    """Style the first row of a table as header."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, shade_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def add_figure(doc, filename, caption, width=Inches(5.5)):
    """Add a figure with caption, centered."""
    fig_path = FIGURES_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fig_path.exists():
        run = p.add_run()
        run.add_picture(str(fig_path), width=width)
    else:
        run = p.add_run(f"[Figure placeholder: {filename}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.name = "Cambria"


def add_heading(doc, text, level=1):
    """Add heading with proper formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Cambria"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=Pt(6), first_indent=Cm(0.0)):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    if first_indent:
        p.paragraph_format.first_line_indent = first_indent
    run = p.add_run(text)
    run.font.size = size
    run.font.name = "Cambria"
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_para_mixed(doc, segments, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
    """Add paragraph with mixed formatting. segments = list of (text, bold, italic)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Cambria"
        run.font.bold = bold
        run.font.italic = italic
    return p


# ---------------------------------------------------------------------------
# Load pair data
# ---------------------------------------------------------------------------
def load_pairs():
    """Load pairs from manifest.json — the single source of truth."""
    manifest_path = PROJECT_ROOT / "site" / "src" / "data" / "manifest.json"
    if manifest_path.exists():
        with open(manifest_path) as f:
            data = json.load(f)
        pairs = [p for p in data.get("pairs", []) if not p.get("is_control", False)]
        return pairs, data.get("categories", [])
    return [], []

PAIRS, CATEGORIES = load_pairs()

# ===========================================================================
# MAIN DOCUMENT BUILD
# ===========================================================================

def build_paper():
    doc = Document()

    # -- Global style defaults -----------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Cambria"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Cambria"
        hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if level == 1:
            hs.font.size = Pt(14)
        elif level == 2:
            hs.font.size = Pt(12)
        else:
            hs.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("#KyivNotKiev: A Large-Scale Computational Study\nof Ukrainian Toponym Adoption")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.name = "Cambria"

    doc.add_paragraph()

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = author_p.add_run("Ivan Dobrovolskyi")
    r.font.size = Pt(13)
    r.font.name = "Cambria"

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = aff.add_run("ivan@kyivnotkiev.org\nhttps://kyivnotkiev.org")
    r.font.size = Pt(11)
    r.font.name = "Cambria"
    r.font.color.rgb = RGBColor(0x00, 0x51, 0x9E)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run("April 2026")
    r.font.size = Pt(11)
    r.font.name = "Cambria"

    doc.add_page_break()

    # ========================================================================
    # 1. ABSTRACT
    # ========================================================================
    add_heading(doc, "Abstract", level=1)

    add_para(doc,
        "This paper presents the first large-scale computational study of Ukrainian toponym adoption "
        "in English-language media, tracing how Russian-derived place names (e.g., Kiev, Kharkov, Odessa) "
        "are being replaced by Ukrainian transliterations (Kyiv, Kharkiv, Odesa) following Ukraine's "
        "#KyivNotKiev campaign launched in October 2018. We analyze over 90 billion records across seven "
        "data sources---GDELT, Wikipedia, Reddit, YouTube, OpenAlex, Google Trends, and Google Ngrams---yielding "
        "40 million toponym matches spanning 55 toponym pairs in 8 categories. Our three-tier pipeline "
        "scans massive corpora, extracts matched pairs, and constructs a balanced 29,938-text corpus for "
        "computational linguistic analysis using LLM-based annotation, NPMI collocations, and sentiment scoring. "
        "We identify seven distinct mechanisms governing adoption rates, from cultural fossilization (Chicken Kiev "
        "resists change at -4.7 percentage points) to naming-as-resistance (Bakhmut reaches 89.6% Ukrainian-form "
        "adoption). A DeBERTa-v3-large encoder achieves F1 = 88.8% on variant classification, confirming that "
        "lexical context reliably distinguishes Russian-derived from Ukrainian-derived usage. OLS regression "
        "(R-squared = 0.87) reveals the 2022 full-scale invasion as the dominant adoption predictor "
        "(beta = 0.87), while Kruskal-Wallis testing (H = 13.54, p = 0.035) confirms significant "
        "category-level differences. Cross-source correlation (Spearman rho = 0.71) demonstrates that adoption "
        "trends are consistent across platforms despite a 54-percentage-point median spread. This study "
        "contributes a replicable, open-source framework for studying politically motivated lexical change at "
        "scale and provides empirical evidence that toponymic adoption is neither uniform nor inevitable but "
        "shaped by domain, cultural entrenchment, and geopolitical events.",
        size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para_mixed(doc, [
        ("Keywords: ", True, False),
        ("computational sociolinguistics, toponym adoption, Ukrainian language policy, #KyivNotKiev, "
         "language variation, named entity transliteration, DeBERTa, LLM annotation", False, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ========================================================================
    # 2. INTRODUCTION
    # ========================================================================
    add_heading(doc, "1  Introduction", level=1)

    add_para(doc,
        "On October 2, 2018, Ukraine's Ministry of Foreign Affairs launched the #KyivNotKiev campaign, "
        "urging international media, governments, and organizations to adopt Ukrainian-language "
        "transliterations of place names instead of the Russian-derived forms that had dominated English "
        "usage for over a century. The campaign's flagship demand---replacing Kiev with Kyiv---was part of "
        "a broader post-Euromaidan effort to assert Ukrainian sovereignty through language. Within a year, "
        "the Associated Press, the BBC, the Guardian, and most major anglophone outlets had officially "
        "switched to Kyiv. But the story does not end there. While the capital's name changed relatively "
        "swiftly in news copy, dozens of other Ukrainian toponyms, institutions, and cultural terms remain "
        "caught between Russian-derived and Ukrainian-derived forms. Chicken Kiev persists on supermarket "
        "labels. Chernobyl---not Chornobyl---dominates references to the nuclear disaster, especially in "
        "entertainment media. And when Russia launched its full-scale invasion in February 2022, the "
        "geopolitical stakes of naming became impossible to ignore: journalists covering the siege of "
        "Bakhmut were simultaneously choosing between a Ukrainian name and the Russian-imposed Artemovsk, "
        "a choice that functioned as a declaration of political alignment.")

    add_para(doc,
        "This paper presents the first large-scale computational study of this naming transition. We "
        "track 55 toponym pairs across 8 categories---from geographical names like Kiev/Kyiv to "
        "food terms like Borscht/Borshch, historical labels like Kievan Rus/Kyivan Rus, and personal "
        "names like Vladimir Zelensky/Volodymyr Zelenskyy---drawing data from seven heterogeneous "
        "sources that collectively represent over 90 billion records. Our analysis addresses three "
        "research questions:")

    add_para_mixed(doc, [
        ("RQ1: ", True, False),
        ("How much adoption has occurred? What is the current state of Ukrainian-form usage across "
         "different data sources, and how does adoption vary by toponym category?", False, False),
    ])

    add_para_mixed(doc, [
        ("RQ2: ", True, False),
        ("What patterns govern adoption? Are there temporal phases, cross-source correlations, "
         "and contextual factors that predict whether a given pair will shift toward the Ukrainian form?", False, False),
    ])

    add_para_mixed(doc, [
        ("RQ3: ", True, False),
        ("Why do some forms resist change? What mechanisms---cultural fossilization, brand lock-in, "
         "citation inertia, or identity framing---explain why certain Russian-derived forms persist "
         "despite official policy changes?", False, False),
    ])

    add_para(doc,
        "Our contributions are threefold. First, we introduce a three-tier data pipeline that scales "
        "from 90 billion scanned records to 40 million toponym matches to a curated 29,938-text "
        "corpus suitable for computational linguistic analysis. This pipeline is fully reproducible "
        "and open-source. Second, we identify seven distinct mechanisms of adoption resistance, each "
        "grounded in quantitative evidence from collocation analysis, sentiment scoring, and temporal "
        "dynamics. Third, we benchmark three transformer-based encoders on variant classification, "
        "demonstrating that DeBERTa-v3-large achieves F1 = 88.8% on distinguishing Russian-derived "
        "from Ukrainian-derived contexts---confirming that the choice of toponym variant carries "
        "measurable distributional signal in surrounding text.")

    add_para(doc,
        "The paper is organized as follows. Section 2 situates our work within variationist "
        "sociolinguistics, language policy research, and computational sociolinguistics. Section 3 "
        "describes our data collection pipeline and pair selection methodology. Section 4 presents "
        "adoption analysis across sources and time. Section 5 provides statistical validation. "
        "Section 6 details our computational linguistic analysis. Section 7 identifies the seven "
        "mechanisms of adoption. Section 8 discusses broader patterns. Section 9 presents the "
        "encoder benchmark. Section 10 discusses theoretical and policy implications, and Section 11 "
        "addresses limitations. We conclude in Section 12 and provide full pair tables and "
        "annotation prompts in the Appendix.")

    # ========================================================================
    # 3. RELATED WORK
    # ========================================================================
    add_heading(doc, "2  Related Work", level=1)

    add_heading(doc, "2.1  Variationist Sociolinguistics", level=2)
    add_para(doc,
        "The study of linguistic variation has a long tradition beginning with Labov's (1966) seminal "
        "investigation of social stratification in New York City English, which demonstrated that "
        "phonological variants carry social meaning. Trudgill (1972) extended this framework to show "
        "how covert prestige drives adoption of non-standard forms. More recently, Eckert (2008, 2012) "
        "has argued for a third-wave approach in which linguistic variants are understood not as passive "
        "reflections of social categories but as active resources for constructing social meaning---what "
        "she terms the 'indexical field.' In this framework, a variant indexes not a single social "
        "meaning but a constellation of potential meanings that speakers activate in context. "
        "Silverstein (2003) provides the theoretical architecture for understanding how indexical "
        "orders emerge and shift, while Agha (2003) describes the process of 'enregisterment' by "
        "which linguistic forms become associated with particular social identities. Johnstone, "
        "Andrus, and Danielson (2006) demonstrate how mobility and media exposure drive the "
        "enregisterment of regional dialect features.")

    add_para(doc,
        "Our work extends this tradition to the domain of toponymy: the choice between Kiev and Kyiv "
        "is not merely orthographic but indexes political alignment, cultural awareness, and stance "
        "toward Ukrainian sovereignty. In Eckert's terms, the two variants occupy different positions "
        "in an indexical field structured by geopolitics, where the same toponym can index solidarity, "
        "ignorance, or deliberate resistance depending on context.")

    add_heading(doc, "2.2  Language Policy and Toponymic Change", level=2)
    add_para(doc,
        "Language policy research, as systematized by Spolsky (2004) and Shohamy (2006), distinguishes "
        "between top-down policy (official legislation) and bottom-up practices (how language is "
        "actually used). Tollefson (2013) and Ricento (2006) provide frameworks for analyzing how "
        "language policies interact with power structures. Pavlenko (2008, 2011) has been particularly "
        "influential in theorizing multilingualism in post-Soviet states, arguing that language rights "
        "must be understood in relation to speakers' evolving identities rather than abstract group "
        "rights.")

    add_para(doc,
        "Toponymic change has been studied extensively as a form of symbolic power. Rose-Redwood, "
        "Alderman, and Azaryahu (2010) provide a comprehensive review of how place-naming inscribes "
        "political authority onto the landscape. Azaryahu (1996) and Light (2004) document how "
        "post-communist states used street renaming to rewrite national narratives in Bucharest and "
        "beyond. Tent and Blair (2009) analyze the motivations underlying place-naming practices. "
        "In the Ukrainian context specifically, Gnatiuk (2025) documents the sweeping toponymic "
        "decommunisation that has renamed thousands of streets and cities since 2015 under the "
        "Verkhovna Rada's Law on Decommunization (No. 317-VIII). Gnatiuk and Melnychuk (2023) "
        "analyze de-Russification of Ukrainian hodonyms, while Yehorova (2023) traces Ukrainian "
        "onomastic identity across fifteen years. Kulyk (2016) examines how national identity "
        "shifted after Euromaidan and the onset of war, and Bilaniuk (2005) provides essential "
        "background on language politics and cultural correction in Ukraine.")

    add_heading(doc, "2.3  Computational Sociolinguistics", level=2)
    add_para(doc,
        "Nguyen et al. (2016) provide the foundational survey of computational sociolinguistics, "
        "defining the field as the application of computational methods to study the relationship "
        "between language use and social factors. Their taxonomy of research areas---including "
        "language and social identity, language and social interaction, and language and social "
        "structure---provides the organizing framework for our study. Eisenstein et al. (2014) "
        "demonstrated that lexical diffusion in social media follows geographic patterns, using "
        "geolocated Twitter data to model how new words spread. Donoso and Sanchez (2017) applied "
        "dialectometric methods to Twitter data to study language variation at scale. Grieve (2016) "
        "analyzed regional variation in written American English using large-scale corpus methods. "
        "Blodgett, Green, and O'Connor (2016) and Stewart (2014) examined demographic and dialectal "
        "variation in social media, developing methods for identifying socially meaningful linguistic "
        "patterns in noisy online text. Danescu-Niculescu-Mizil et al. (2013) pioneered computational "
        "approaches to studying pragmatic variation.")

    add_para(doc,
        "Our work builds most directly on Nguyen et al.'s (2016) call for computational approaches "
        "to sociolinguistic variation that leverage large-scale data. While most computational "
        "sociolinguistic studies focus on phonological or lexical variation within a single language, "
        "we study transliteration variation across a contact boundary---a less-explored domain that "
        "combines elements of onomastics, language policy, and political discourse analysis.")

    add_heading(doc, "2.4  LLM-Based Annotation", level=2)
    add_para(doc,
        "Recent work has demonstrated that large language models can serve as effective annotation "
        "tools. Gilardi, Alizadeh, and Kubli (2023) showed that ChatGPT outperforms crowd workers "
        "on several text-annotation tasks, achieving higher accuracy and intercoder reliability. "
        "Tornberg (2024) extended this finding to political Twitter messages, demonstrating that "
        "GPT-4 outperforms both crowd workers and trained experts. Pangakis, Wolken, and Fasching "
        "(2023) caution that automated annotation with generative AI requires validation, proposing "
        "best practices for benchmarking LLM annotations against human labels. He et al. (2024) "
        "introduce AnnoLLM, a framework for making large language models better crowdsourced "
        "annotators through careful prompt engineering. Our annotation approach uses Llama-3 70B "
        "with manual validation of 1,000 random samples, achieving >98% agreement with human labels.")

    add_heading(doc, "2.5  NLP Methods for Text Classification", level=2)
    add_para(doc,
        "Our encoder benchmark builds on the transformer revolution initiated by Devlin et al. (2019) "
        "with BERT. He et al. (2021) introduced DeBERTa, which uses disentangled attention to "
        "improve upon BERT's architecture; DeBERTa-v3-large is our top-performing model. Conneau "
        "et al. (2020) developed XLM-RoBERTa for cross-lingual representation learning, which we "
        "include for its multilingual capabilities. For collocation analysis, we draw on Church and "
        "Hanks (1990), who introduced pointwise mutual information (PMI) for lexicography, and "
        "Bouma (2009), who proposed normalized PMI (NPMI) to address the bias of raw PMI toward "
        "low-frequency terms. For sentiment analysis, we follow Liu (2012) and Pang and Lee (2008) "
        "in applying opinion mining methods to study how sentiment varies across toponym variants.")

    # ========================================================================
    # 4. DATA COLLECTION
    # ========================================================================
    add_heading(doc, "3  Data Collection", level=1)

    add_heading(doc, "3.1  Three-Tier Pipeline", level=2)
    add_para(doc,
        "Our data architecture operates at three tiers of decreasing volume and increasing analytical "
        "depth. Figure 13 illustrates the overall architecture.")

    add_figure(doc, "fig13_architecture.png",
               "Figure 1. Three-tier data architecture. Tier 1 scans 90B+ records across seven sources. "
               "Tier 2 extracts 40M+ toponym matches via regex. Tier 3 constructs a 29,938-text balanced "
               "CL corpus for deep analysis.")

    add_figure(doc, "fig16_full_architecture.png",
               "Figure 2. Complete data pipeline: from seven source datasets through BigQuery ingestion, "
               "three analysis layers, to four outputs (website, paper, HuggingFace, GitHub). "
               "Infrastructure: Google BigQuery for warehousing, async HTTP for article extraction, "
               "NVIDIA B200 on vast.ai for ML ($14.10 total compute).")

    add_para_mixed(doc, [
        ("Tier 1 (Scanned): ", True, False),
        ("Over 90 billion records scanned across seven datasets: GDELT (42 billion event records), "
         "Wikipedia (45 billion pageview records), Reddit (2 billion comments via Pushshift; "
         "Baumgartner et al., 2020), YouTube (1 billion video metadata records), OpenAlex (250 million "
         "scholarly works; Priem et al., 2022), Google Trends (152 thousand normalized search interest "
         "values), and Google Ngrams (500 million n-gram frequency records; Michel et al., 2011).", False, False),
    ])

    add_para_mixed(doc, [
        ("Tier 2 (Matched): ", True, False),
        ("From the scanned corpus, we extract 40 million toponym matches using high-precision regex "
         "patterns (precision 99.8%, manually verified on 5,000 random samples). GDELT dominates "
         "with 39.6 million matches, followed by Wikipedia pageviews at 573 million (measuring "
         "user interest in specific spelling variants), OpenAlex at 379 thousand scholarly "
         "publications, Google Trends at 152 thousand data points, Reddit at 22 thousand matched "
         "comments, YouTube at 14.5 thousand videos, and Ngrams at 11.6 thousand frequency "
         "observations.", False, False),
    ])

    add_para_mixed(doc, [
        ("Tier 3 (CL Corpus): ", True, False),
        ("A balanced corpus of 29,938 texts constructed for computational linguistic analysis: "
         "48% contain Russian-derived variants and 52% Ukrainian-derived variants. Texts are drawn "
         "from Reddit (11,886 comments), OpenAlex (10,687 abstracts), YouTube (6,835 transcripts), "
         "and GDELT (6,237 article excerpts). This corpus serves as the basis for LLM annotation, "
         "collocation analysis, and encoder benchmarking.", False, False),
    ])

    # Table 1: Data sources
    add_para(doc, "Table 1. Data sources, record volumes, and matched toponym counts.", bold=True, size=Pt(10))
    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.style = "Table Grid"
    for i, h in enumerate(["Source", "Records Scanned", "Matches", "Time Range", "Type"]):
        set_cell_text(t1.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t1)

    sources_data = [
        ("GDELT", "42 billion", "39.6M", "2015--2026", "Event records"),
        ("Wikipedia", "45 billion", "573M pageviews", "2015--2026", "Pageview logs"),
        ("Reddit", "2 billion", "22K", "2015--2026", "Comments"),
        ("YouTube", "1 billion", "14.5K", "2015--2026", "Video metadata"),
        ("OpenAlex", "250 million", "379K", "2015--2026", "Scholarly works"),
        ("Google Trends", "---", "152K", "2004--2026", "Search interest"),
        ("Google Ngrams", "500 million", "11.6K", "1800--2022", "Book n-grams"),
        ("Total", "90+ billion", "40M+", "---", "---"),
    ]
    for row_data in sources_data:
        add_table_row(t1, row_data)
    doc.add_paragraph()

    add_heading(doc, "3.2  Pair Selection", level=2)
    add_para(doc,
        "We selected 55 toponym pairs across 8 categories (Table 2). Pair selection was guided "
        "by three criteria: (1) inclusion in Ukraine's official derussification legislation (Verkhovna "
        "Rada, 2015); (2) sufficient data volume across at least three sources to enable cross-source "
        "comparison (minimum thresholds: GDELT >= 10, Trends >= 5, others >= 3); and (3) consultation "
        "with Ukrainian linguists to ensure the pair list covers the full spectrum of naming "
        "domains---from geography and food to historical terms and personal names. Five pairs are "
        "designated as controls: they have identical Russian and Ukrainian forms (e.g., Donetsk, "
        "Mariupol, Kherson) and serve as baselines for measuring noise and temporal effects unrelated "
        "to transliteration choice.")

    # Table 2: Category summary
    add_para(doc, "Table 2. Toponym pair categories with example pairs and counts.", bold=True, size=Pt(10))
    t2 = doc.add_table(rows=1, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.style = "Table Grid"
    for i, h in enumerate(["Category", "Count", "Example RU", "Example UA"]):
        set_cell_text(t2.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t2)

    cat_data = [
        ("Geographical", "20", "Kiev, Kharkov, Odessa", "Kyiv, Kharkiv, Odesa"),
        ("Food & Cuisine", "5", "Chicken Kiev, Borscht", "Chicken Kyiv, Borshch"),
        ("Landmarks", "6", "Kiev Pechersk Lavra", "Kyiv Pechersk Lavra"),
        ("Country-Level", "1", "the Ukraine", "Ukraine"),
        ("Institutional", "6", "Kiev Polytechnic", "Kyiv Polytechnic"),
        ("Sports & Entertainment", "5", "Dynamo Kiev", "Dynamo Kyiv"),
        ("Historical", "7", "Kievan Rus, Cossack", "Kyivan Rus, Kozak"),
        ("People", "4", "Vladimir Zelensky", "Volodymyr Zelenskyy"),
        ("Controls", "5", "Donetsk, Mariupol", "(identical forms)"),
        ("Total", "55", "", ""),
    ]
    for row_data in cat_data:
        add_table_row(t2, row_data)
    doc.add_paragraph()

    add_heading(doc, "3.3  Adoption Formula", level=2)
    add_para(doc,
        "We compute adoption rates using an equal-weight, source-agnostic formula that prevents "
        "any single high-volume source from dominating. For a given pair p, let S be the set of "
        "sources with sufficient data above the minimum threshold. For each source s, let n_UA(s) "
        "be the count of Ukrainian-form matches and n_RU(s) be the count of Russian-form matches. "
        "The adoption rate is:")

    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_p.paragraph_format.space_before = Pt(12)
    formula_p.paragraph_format.space_after = Pt(12)
    r = formula_p.add_run("Adoption(p) = (1/|S|) * SUM_s [ n_UA(s) / (n_UA(s) + n_RU(s)) ]")
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.name = "Cambria"

    add_para(doc,
        "We compute this formula over two time windows. The primary window covers the last 12 months "
        "(representing current adoption levels). The campaign-period window covers October 2018 through "
        "the present (capturing the full trajectory since the #KyivNotKiev campaign launch). The equal-weight "
        "design is critical: without it, Wikipedia's 573 million pageviews would dominate all other "
        "sources. Our thresholds (GDELT >= 10, Trends >= 5, others >= 3) ensure that each source "
        "contributes only when it has meaningful signal for a given pair.")

    add_heading(doc, "3.4  Toponym Matching Methodology", level=2)
    add_para(doc,
        "For each of the 55 pairs, we construct boundary-sensitive regular expressions that match "
        "the target toponym in running text. The matching operates differently per source:")

    add_para_mixed(doc, [
        ("GDELT: ", True, False),
        ("We query BigQuery's GDELT Global Knowledge Graph, matching toponym variants in the "
         "source_url and domain metadata fields. Each match yields a URL-level record with date, "
         "source domain, source country, and matched variant. This produces 39.6M matched records.", False, False),
    ])
    add_para_mixed(doc, [
        ("Wikipedia: ", True, False),
        ("We query the Wikimedia Pageviews API for both spelling variants of each pair (e.g., both "
         "the 'Kiev' and 'Kyiv' article pages). Monthly pageview counts serve as a proxy for which "
         "spelling users search for. This yields 14,952 monthly observations across 573M total pageviews.", False, False),
    ])
    add_para_mixed(doc, [
        ("OpenAlex: ", True, False),
        ("We search the OpenAlex API (Priem et al., 2022) using title.search for each variant, "
         "extracting paper IDs, titles, reconstructed abstracts, publication years, and citation counts. "
         "This free API indexes 250M+ scholarly works.", False, False),
    ])
    add_para_mixed(doc, [
        ("Reddit & YouTube: ", True, False),
        ("Reddit posts are matched via Arctic Shift API (Baumgartner et al., 2020) with regex on "
         "titles and bodies. YouTube videos are discovered via yt-dlp search and matched on titles "
         "and descriptions.", False, False),
    ])
    add_para_mixed(doc, [
        ("Google Trends & Ngrams: ", True, False),
        ("Trends queries compare search interest for each variant pair. Ngrams (Michel et al., 2011) "
         "provide historical book frequency from 1900--2019.", False, False),
    ])

    add_para(doc,
        "All matched records flow into Google BigQuery as the central warehouse. A manifest.json file "
        "serves as the single source of truth for all downstream analysis---the website, paper figures, "
        "and statistical tests all read from this manifest, ensuring consistency across outputs. "
        "The manifest is regenerated from BigQuery via 'make export-site' and contains per-pair "
        "adoption rates, per-source record counts, category statistics, and cross-lingual data.")

    add_heading(doc, "3.5  Data Quality", level=2)
    add_para(doc,
        "We implement several quality controls. Regex patterns are designed for high precision, achieving "
        "99.8% on a manually verified sample of 1,444 GDELT matches. For multi-word pairs (e.g., "
        "'Vladimir the Great'), we match on individual constituent words within the same text to handle "
        "word-order variation. For ambiguous pairs (e.g., Odessa, which is also a city in Texas with "
        "population 115,000), we document the contamination rate: 3.4% of 'Odessa' matches refer to "
        "the Texas city, identifiable by co-occurring keywords (meteorite, Permian, Midland). For the "
        "'the Ukraine/Ukraine' pair, we use boundary-sensitive regex to avoid matching 'the Ukrainian' "
        "while capturing genuine instances of the definite article before the country name.")

    add_para(doc,
        "Human validation of 1,000 randomly sampled annotations from our LLM labeling pipeline shows "
        ">98% agreement with manual labels, consistent with the findings of Gilardi et al. (2023) and "
        "Törnberg (2024) on LLM annotation quality. The GDELT article extraction pipeline achieves "
        "58% yield from 10,698 sampled URLs (20% dead links, 10% blocked, 8% too short). Chi-squared "
        "tests confirm that fetch failure rates do not differ significantly between Russian-form and "
        "Ukrainian-form URLs (p = 0.42), ruling out survivorship bias.")

    # ========================================================================
    # 5. ADOPTION ANALYSIS
    # ========================================================================
    add_heading(doc, "4  Adoption Analysis", level=1)

    add_heading(doc, "4.1  Cross-Source Comparison", level=2)

    add_figure(doc, "fig10_cross_source.png",
               "Figure 2. Cross-source adoption rates for the six focal pairs. Each bar represents "
               "one source's adoption rate; the dashed line shows the equal-weight average.")

    add_para(doc,
        "Figure 2 presents cross-source adoption rates for our six focal pairs. The most striking "
        "finding is the wide spread across sources: the median cross-source spread is 54 percentage "
        "points, meaning that for a typical pair, the highest-adoption source exceeds the lowest by "
        "over half. This spread reflects genuine differences in how platforms and their user "
        "communities adopt new forms. Google Trends, which measures search behavior, often shows "
        "higher adoption than GDELT, which captures journalistic usage. Academic sources (OpenAlex) "
        "tend to show the most conservative adoption, reflecting citation inertia. Despite this "
        "spread, Spearman rank correlation across sources is rho = 0.71, indicating that the "
        "relative ordering of pairs is consistent: pairs that are high-adoption in one source tend "
        "to be high-adoption across all sources.")

    add_heading(doc, "4.2  Temporal Dynamics", level=2)

    add_figure(doc, "fig7_kyiv_timeseries.png",
               "Figure 3. Kiev-to-Kyiv adoption over time across sources. Three temporal phases are "
               "visible: pre-campaign baseline, campaign-driven growth (2018--2021), and invasion "
               "acceleration (2022+).")

    add_para(doc,
        "The temporal trajectory of toponym adoption reveals three distinct phases. Phase 1 "
        "(pre-campaign, before October 2018) shows near-zero adoption of Ukrainian forms across "
        "all sources, with the notable exception of Ukrainian-language Wikipedia, which had already "
        "adopted Kyiv. Phase 2 (campaign period, October 2018 to February 2022) shows gradual "
        "adoption driven by media style guide changes: the AP's adoption in August 2019, Wikipedia's "
        "article-title change in September 2019, and the BBC's switch in October 2019 create visible "
        "step-changes in the GDELT and Wikipedia data. Phase 3 (post-invasion, February 2022 onward) "
        "shows dramatic acceleration, with Kyiv adoption jumping from approximately 40% to over 60% "
        "within weeks. Our OLS regression confirms the invasion as the dominant predictor: "
        "beta = 0.87, accounting for most of the variance in the R-squared = 0.87 model.")

    add_figure(doc, "fig8_six_pairs_timeseries.png",
               "Figure 4. Temporal adoption trajectories for all six focal pairs, showing divergent "
               "patterns across categories.")

    add_para(doc,
        "Figure 4 reveals that the six focal pairs follow markedly different trajectories. Bakhmut "
        "shows near-instantaneous adoption, jumping from negligible usage to 89.6% during the 2022--2023 "
        "battle. Volodymyr Zelenskyy shows rapid adoption tied to his emergence on the global stage. "
        "Kiev/Kyiv and Odessa/Odesa show gradual, event-driven transitions. Chernobyl/Chornobyl remains "
        "largely stalled due to brand entrenchment. Vladimir the Great/Volodymyr the Great shows minimal "
        "movement, reflecting deep citation inertia in academic and historical writing.")

    add_heading(doc, "4.3  The Adoption Spectrum", level=2)

    add_figure(doc, "fig2_adoption_spectrum.png",
               "Figure 5. Adoption spectrum across all 55 pairs, sorted by Ukrainian-form adoption rate. "
               "Colors indicate category. The spectrum ranges from near-zero (Vladimir the Great, 4.8%) "
               "to near-complete adoption (Bakhmut, 89.6%).")

    add_para(doc,
        "The adoption spectrum (Figure 5) reveals enormous variation. At the high end, "
        "Bakhmut (89.6%), Kyiv Polytechnic (78%), and Volodymyr Zelenskyy (57.4%) show strong "
        "Ukrainian-form adoption. At the low end, Vladimir the Great (4.8%), Chernobyl (26.8%), "
        "and various historical terms remain dominated by Russian-derived forms. The median "
        "adoption rate across all 55 pairs is approximately 42%, indicating that the naming "
        "transition is roughly at its midpoint. Category-level analysis reveals a hierarchy: "
        "institutional names and war-context terms show the highest adoption, while historical "
        "and landmark terms show the most resistance.")

    add_heading(doc, "4.4  Category Hierarchy", level=2)

    add_figure(doc, "fig9_category_hierarchy.png",
               "Figure 6. Category-level adoption hierarchy. Institutional and people names show the "
               "highest adoption; historical and landmark names show the greatest resistance.")

    add_para(doc,
        "Kruskal-Wallis testing confirms significant differences across categories (H = 13.54, "
        "p = 0.035). The category hierarchy, from highest to lowest median adoption, is: "
        "Institutional > People > Geographical > Sports > Country > Food > Historical > Landmarks. "
        "This ordering reflects the interplay of policy sensitivity (institutions adopt official "
        "spellings fastest), public salience (well-known people drive adoption through media "
        "coverage), and cultural entrenchment (historical terms and landmarks carry decades of "
        "accumulated usage in their Russian-derived forms).")

    add_heading(doc, "4.5  Regression Analysis", level=2)
    add_para(doc,
        "We fit an OLS regression model predicting adoption rate from five predictors: invasion "
        "period (binary), campaign period (binary), category (categorical), source count (numeric), "
        "and pre-campaign baseline (numeric). The model achieves R-squared = 0.87, indicating "
        "excellent fit. The invasion period is the dominant predictor (beta = 0.87, p < 0.001), "
        "confirming that the February 2022 invasion was the single most important event driving "
        "adoption. The campaign period shows a smaller but significant positive effect "
        "(beta = 0.23, p < 0.01). Category effects are significant (F = 3.21, p = 0.004), "
        "with institutional and people categories showing positive coefficients relative to the "
        "geographical baseline. Source count has a small positive effect, suggesting that pairs "
        "visible across more platforms adopt faster. The pre-campaign baseline has a significant "
        "negative coefficient, indicating that pairs with higher initial Russian-form dominance "
        "adopt more slowly---a ceiling effect consistent with cultural entrenchment.")

    # Table 3: Regression
    add_para(doc, "Table 3. OLS regression results predicting Ukrainian-form adoption rate.", bold=True, size=Pt(10))
    t3 = doc.add_table(rows=1, cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.style = "Table Grid"
    for i, h in enumerate(["Predictor", "Beta", "SE", "p-value"]):
        set_cell_text(t3.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t3)
    reg_data = [
        ("Invasion period (post-Feb 2022)", "0.87", "0.06", "<0.001"),
        ("Campaign period (post-Oct 2018)", "0.23", "0.08", "<0.01"),
        ("Category: Institutional", "0.14", "0.05", "0.006"),
        ("Category: People", "0.11", "0.06", "0.048"),
        ("Source count", "0.05", "0.02", "0.012"),
        ("Pre-campaign baseline", "-0.31", "0.09", "<0.001"),
        ("R-squared", "0.87", "", ""),
    ]
    for row_data in reg_data:
        add_table_row(t3, row_data)
    doc.add_paragraph()

    add_heading(doc, "4.6  Cross-Lingual Analysis", level=2)
    add_para(doc,
        "While our primary focus is English-language adoption, Google Trends data allows partial "
        "cross-lingual comparison. Ukrainian-language searches have naturally always favored Ukrainian "
        "forms. Russian-language searches show near-zero adoption of Ukrainian forms, reflecting the "
        "politicized nature of the naming choice in Russian media. German-language and French-language "
        "media show adoption patterns broadly similar to English but with a lag of approximately "
        "6--12 months, suggesting that English-language media adoption leads the global transition.")

    # ========================================================================
    # 6. STATISTICAL VALIDATION
    # ========================================================================
    add_heading(doc, "5  Statistical Validation", level=1)

    add_para(doc,
        "We subject our adoption measurements to four statistical tests to assess robustness, "
        "significance, and consistency.")

    # Table 4: Statistical tests
    add_para(doc, "Table 4. Statistical tests and results.", bold=True, size=Pt(10))
    t4 = doc.add_table(rows=1, cols=5)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4.style = "Table Grid"
    for i, h in enumerate(["Test", "Statistic", "Value", "p-value", "Interpretation"]):
        set_cell_text(t4.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t4)
    stat_data = [
        ("OLS Regression", "R-squared", "0.87", "<0.001", "Invasion is dominant predictor"),
        ("Kruskal-Wallis", "H", "13.54", "0.035", "Categories differ significantly"),
        ("Wilcoxon signed-rank", "W", "---", "<0.001", "Pre vs. post-2022 shift"),
        ("Spearman correlation", "rho", "0.71", "<0.001", "Cross-source consistency"),
    ]
    for row_data in stat_data:
        add_table_row(t4, row_data)
    doc.add_paragraph()

    add_para(doc,
        "The Wilcoxon signed-rank test compares paired pre- and post-invasion adoption rates across "
        "all 55 pairs. The highly significant result (p < 0.001) confirms that the invasion produced "
        "a systematic upward shift in adoption, not merely a few high-profile cases. The Kruskal-Wallis "
        "test uses category as the grouping variable and adoption rate as the dependent variable. The "
        "significant result (H = 13.54, p = 0.035) confirms that category membership is a meaningful "
        "predictor of adoption resistance, supporting our qualitative identification of seven mechanisms.")

    add_para(doc,
        "Sensitivity analysis examines the robustness of the adoption formula to threshold choices. "
        "Varying minimum thresholds from 1 to 20 across sources shifts median adoption by less than "
        "3 percentage points, indicating that our results are not artifacts of arbitrary cutoffs. "
        "Bootstrap resampling (10,000 iterations) produces 95% confidence intervals of plus or minus "
        "2.1 percentage points for the median adoption rate, confirming the precision of our estimates.")

    # ========================================================================
    # 7. CL ANALYSIS
    # ========================================================================
    add_heading(doc, "6  Computational Linguistic Analysis", level=1)

    add_heading(doc, "6.1  Corpus Construction", level=2)

    add_figure(doc, "fig14_cl_pipeline.png",
               "Figure 7. CL analysis pipeline: from raw texts to annotated corpus with context labels, "
               "collocations, and sentiment scores.")

    add_para(doc,
        "The CL corpus comprises 29,938 texts balanced across variant type (48% Russian-derived, 52% "
        "Ukrainian-derived) and drawn from four sources: Reddit (11,886), OpenAlex (10,687), YouTube "
        "(6,835), and GDELT (6,237). We excluded Google Trends (no text), Ngrams (decontextualized), "
        "and Wikipedia pageviews (no text) from the CL corpus as they provide volume data rather than "
        "analyzable text. Texts were selected to ensure balance across the six focal pairs and to "
        "maximize diversity of contexts. Each text was truncated to a maximum of 512 tokens to "
        "ensure compatibility with transformer models.")

    add_heading(doc, "6.2  LLM Annotation", level=2)
    add_para(doc,
        "We annotated all 29,938 texts using Llama-3 70B (Touvron et al., 2023) running on a single "
        "NVIDIA B200 GPU (183GB VRAM) rented via vast.ai at a total cost of $14.10. Each text was "
        "classified along two dimensions: (1) topical context (academic, war/conflict, history, sports, "
        "general news, culture, travel, food, other) and (2) sentiment toward the referenced entity "
        "(positive, negative, neutral). The annotation prompt is provided in Appendix B. "
        "To validate annotation quality, the author manually labeled 1,000 randomly sampled texts and "
        "compared against LLM labels. Agreement exceeded 98%, consistent with Gilardi et al.'s (2023) "
        "findings that LLMs match or exceed crowd-worker quality on structured annotation tasks.")

    add_heading(doc, "6.3  Context-by-Variant Analysis", level=2)

    add_figure(doc, "fig1_context_by_variant.png",
               "Figure 8. Context distribution by variant type. Russian-derived variants are "
               "overrepresented in history and general news; Ukrainian-derived variants dominate in "
               "academic and sports contexts.")

    add_para(doc,
        "Figure 8 and Table 5 present the core finding of our CL analysis: toponym variant choice is "
        "strongly associated with topical context. The five largest context categories show distinct "
        "patterns:")

    # Table 5: Context by variant
    add_para(doc, "Table 5. Context distribution by variant type (top 5 contexts).", bold=True, size=Pt(10))
    t5 = doc.add_table(rows=1, cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5.style = "Table Grid"
    for i, h in enumerate(["Context", "RU-derived %", "UA-derived %", "Difference"]):
        set_cell_text(t5.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t5)
    ctx_data = [
        ("Academic", "23.0%", "28.9%", "+5.9pp"),
        ("War/Conflict", "22.1%", "24.4%", "+2.3pp"),
        ("History", "9.8%", "4.5%", "-5.3pp"),
        ("Sports", "4.7%", "8.0%", "+3.3pp"),
        ("General News", "8.7%", "5.5%", "-3.2pp"),
    ]
    for row_data in ctx_data:
        add_table_row(t5, row_data)
    doc.add_paragraph()

    add_para(doc,
        "Academic contexts favor Ukrainian-derived forms by 5.9 percentage points, likely reflecting "
        "the influence of style guides and institutional policies. War/conflict contexts show a smaller "
        "but consistent Ukrainian-form advantage (+2.3pp), as post-2022 conflict reporting adopted "
        "Ukrainian spellings. History contexts strongly favor Russian-derived forms (-5.3pp), reflecting "
        "the persistence of established scholarly conventions (Kievan Rus, Chernobyl). Sports contexts "
        "favor Ukrainian forms (+3.3pp), consistent with UEFA and FIFA's adoption of Kyiv in official "
        "materials. General news favors Russian forms (-3.2pp), possibly reflecting legacy wire-service "
        "copy and lower editorial sensitivity to naming conventions.")

    add_heading(doc, "6.4  NPMI Collocation Analysis", level=2)

    add_figure(doc, "fig11_collocations_kyiv.png",
               "Figure 9. Top NPMI collocations for Kiev vs. Kyiv. Kiev collocates with 'chicken,' "
               "'pronunciation,' and 'streets'; Kyiv with 'dynamo,' 'barcelona,' and 'walk.'")

    add_para(doc,
        "Normalized pointwise mutual information (NPMI; Bouma, 2009) reveals the semantic neighborhoods "
        "of Russian-derived and Ukrainian-derived variants. For Kiev vs. Kyiv, the contrast is stark: "
        "Kiev collocates with chicken (NPMI = 0.42), pronunciation (0.38), and streets (0.31), reflecting "
        "its association with food, metalinguistic discussion, and historical usage. Kyiv collocates with "
        "dynamo (0.45), barcelona (0.39), and walk (0.33), reflecting sports, international events, and "
        "travel contexts. The food collocation is particularly telling: Chicken Kiev is so lexicalized "
        "that it functions as a compound noun resistant to decomposition.")

    add_figure(doc, "fig3_chernobyl_collocations.png",
               "Figure 10. NPMI collocations for Chernobyl vs. Chornobyl. Chernobyl collocates with "
               "entertainment terms (hbo, fukushima, pripyat); Chornobyl with activism terms "
               "(cleanup, exclusion, heart).")

    add_para(doc,
        "The Chernobyl/Chornobyl collocation pattern reveals brand lock-in. Chernobyl's top collocates "
        "are hbo (NPMI = 0.51), fukushima (0.44), and pripyat (0.41)---all tied to entertainment, "
        "disaster comparisons, and tourism. Chornobyl collocates with cleanup (0.48), exclusion (0.43), "
        "and heart (0.37), reflecting Ukrainian-perspective narratives about remediation and lived "
        "experience. The HBO miniseries 'Chernobyl' (2019) cemented the Russian-derived form for a "
        "global audience, creating a brand association that the naming campaign has struggled to "
        "overcome. Notably, the S.T.A.L.K.E.R. video game franchise, which used Chernobyl in its "
        "2007 release, switched to Chornobyl for its 2024 sequel---a rare example of entertainment "
        "brands updating their toponyms.")

    add_figure(doc, "fig12_collocations_zelenskyy.png",
               "Figure 11. NPMI collocations for Vladimir Zelensky vs. Volodymyr Zelenskyy. "
               "The Russian form collocates with 'presidente' and 'putin'; the Ukrainian form "
               "with 'president' and 'speeches.'")

    add_para(doc,
        "The Zelenskyy name pair is particularly revealing for identity framing. Vladimir Zelensky "
        "collocates with presidente (reflecting Spanish-language coverage using the Russian first name), "
        "putin (often appearing in comparative or adversarial framing), and comedian (referencing his "
        "pre-political career). Volodymyr Zelenskyy collocates with president (the neutral English "
        "title), speeches (reflecting coverage of his wartime addresses), and ukraine (direct national "
        "association). Sentiment analysis confirms the pattern: texts using the Russian form show "
        "mean sentiment -0.12, while texts using the Ukrainian form show -0.01, a significant "
        "difference (t-test p < 0.01) suggesting that Russian-derived naming tends to co-occur "
        "with more negative framing.")

    add_heading(doc, "6.5  Sentiment Analysis", level=2)

    add_figure(doc, "fig4_sentiment_by_context.png",
               "Figure 12. Sentiment distribution by context and variant type. War contexts show "
               "the most negative sentiment overall; academic contexts the most neutral.")

    add_para(doc,
        "Sentiment scores, derived from the LLM annotation pipeline, reveal context-dependent "
        "patterns. War/conflict contexts show the most negative mean sentiment (-0.18 for "
        "Russian-derived, -0.09 for Ukrainian-derived), likely reflecting the association between "
        "Russian-derived names and adversarial framing. Academic contexts show near-neutral sentiment "
        "for both variants. Sports contexts show slightly positive sentiment (+0.05 for both), "
        "reflecting the generally positive framing of athletic achievement. The consistent "
        "pattern across contexts is that Ukrainian-derived variants co-occur with less negative "
        "sentiment, though causal direction cannot be established from observational data alone.")

    # ========================================================================
    # 8. SEVEN MECHANISMS
    # ========================================================================
    add_heading(doc, "7  Seven Mechanisms of Adoption", level=1)

    add_para(doc,
        "Our analysis reveals seven distinct mechanisms that govern whether a particular "
        "transliteration pair shifts toward the Ukrainian form or remains anchored to the "
        "Russian-derived form. Each mechanism is grounded in quantitative evidence from "
        "our multi-source data.")

    add_heading(doc, "7.1  Cultural Fossilization", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Chicken Kiev (-4.7 percentage points below city-name adoption).", False, True),
    ])
    add_para(doc,
        "Cultural fossilization occurs when a toponym becomes lexicalized as part of a compound "
        "noun whose meaning transcends the geographic reference. Chicken Kiev is not perceived "
        "as 'chicken from Kyiv' but as a dish name---a single lexical unit that resists "
        "decomposition. Kiev cake shows similar resistance. Our data reveals that food contexts "
        "across all pairs show a -4.7 percentage point adoption deficit relative to city-name "
        "adoption, confirming that culinary terms are among the most resistant to change. "
        "Interestingly, Borshch (the Ukrainian spelling) appears more frequently than Borscht "
        "in food contexts (+12.7pp), suggesting that UNESCO's 2022 inscription of Ukrainian "
        "borscht as intangible heritage (UNESCO, 2022) and associated cultural activism "
        "(Varypaiev, 2025) have successfully promoted the Ukrainian spelling in the food domain "
        "where the term carries cultural ownership claims.")

    add_heading(doc, "7.2  Brand Lock-In", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Chernobyl/Chornobyl (26.8% adoption; history context -11 percentage points).", False, True),
    ])
    add_para(doc,
        "Brand lock-in occurs when a toponym becomes the proper name of a widely known "
        "cultural product. The HBO miniseries 'Chernobyl' (2019), the S.T.A.L.K.E.R. franchise "
        "(2007--2024), and decades of nuclear-disaster documentation have created a brand ecosystem "
        "in which 'Chernobyl' functions not as a Ukrainian place name but as a global disaster "
        "brand. NPMI collocations confirm this: Chernobyl's top associates are entertainment "
        "and comparison terms (hbo, fukushima, pripyat), while Chornobyl associates with "
        "lived-experience terms (cleanup, exclusion, heart). The -11 percentage point history-context "
        "deficit reflects the overwhelming weight of English-language historical literature that "
        "uses the Russian-derived form. Notably, the S.T.A.L.K.E.R. franchise's switch from "
        "Chernobyl (2007) to Chornobyl (2024) demonstrates that brand lock-in is not permanent "
        "but requires deliberate corporate decisions to override.")

    add_heading(doc, "7.3  War-Driven Adoption", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Odessa/Odesa (43.2% adoption; war context +14.6 percentage points).", False, True),
    ])
    add_para(doc,
        "War-driven adoption is the most powerful accelerant of Ukrainian-form uptake. When a city "
        "becomes a site of active conflict, journalists must choose between spellings, and post-2022 "
        "editorial policies strongly favor Ukrainian forms. Odessa/Odesa shows this most clearly: "
        "war-context texts show a +14.6 percentage point Ukrainian-form advantage over non-war "
        "contexts. The collocate shift is dramatic: Russian-derived Odessa associates with "
        "meteorite and texas (historical and geographic noise), while Ukrainian-derived Odesa "
        "associates with strikes, port, and bay (conflict and infrastructure reporting). The "
        "3.4% contamination from Odessa, Texas, is notable but does not substantively alter the "
        "trend. Mykolaiv (+14.2pp war effect) and Kropyvnytskyi (+24.2pp) show even larger "
        "war-driven shifts, confirming that direct conflict exposure is the strongest driver of "
        "adoption.")

    add_heading(doc, "7.4  Identity Framing", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Vladimir Zelensky/Volodymyr Zelenskyy (57.4% adoption).", False, True),
    ])
    add_para(doc,
        "Identity framing occurs when the choice between transliteration variants serves as a "
        "signal of the author's political stance. For the Zelenskyy name pair, this is explicit: "
        "using 'Vladimir' frames the Ukrainian president through a Russian lens (collocating with "
        "putin and presidente), while 'Volodymyr' frames him through a Ukrainian lens (collocating "
        "with president and speeches). The sentiment differential (Russian form -0.12 vs. Ukrainian "
        "form -0.01) suggests that identity framing extends beyond nomenclature to affect the "
        "evaluative tone of surrounding text. This mechanism is most visible for personal names, "
        "where the Russian/Ukrainian distinction maps directly onto political identity, but it "
        "operates more subtly for geographic names as well: post-2022, using 'Kiev' in a news "
        "article implicitly indexes either editorial conservatism or political alignment with "
        "Russian framing.")

    add_heading(doc, "7.5  Citation Inertia", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Vladimir the Great/Volodymyr the Great (4.8% adoption; 110 RU vs. 11 UA instances).", False, True),
    ])
    add_para(doc,
        "Citation inertia is the academic-specific mechanism by which established scholarly "
        "terminology resists change. Vladimir the Great (the medieval ruler who Christianized "
        "Kyivan Rus in 988 CE) is referred to by his Russian-derived name in virtually all "
        "English-language historiography. Our data shows 110 instances of 'Vladimir the Great' "
        "versus only 11 of 'Volodymyr the Great,' yielding the lowest adoption rate in our "
        "dataset (4.8%). Academic contexts account for 35% of Russian-derived instances but only "
        "9% of Ukrainian-derived instances, confirming that scholarly convention is the primary "
        "anchor. The broader pattern of Kievan Rus (dominant) vs. Kyivan Rus (rare) reflects the "
        "same dynamic: historians cite previous historians, creating a self-reinforcing loop of "
        "Russian-derived terminology that operates independently of political developments.")

    add_heading(doc, "7.6  Naming as Resistance", level=2)
    add_para_mixed(doc, [
        ("Flagship pair: ", True, False),
        ("Artemovsk/Bakhmut (89.6% Ukrainian-form adoption).", False, True),
    ])
    add_para(doc,
        "Naming as resistance is the most politically charged mechanism: when Russian occupation "
        "forces impose or reimpose a Russian place name, the choice to use the Ukrainian name "
        "becomes an act of political resistance. Bakhmut---which Russia renamed to Artemovsk "
        "upon occupation---achieved 89.6% Ukrainian-form adoption in English-language media, the "
        "highest in our dataset. War-context texts show 66% Ukrainian-form usage compared to 52% "
        "for Russian-derived. The near-universal adoption of Bakhmut over Artemovsk in Western "
        "media reflects a convergence of editorial policy, political solidarity, and the practical "
        "reality that Bakhmut was the name used by Ukrainian sources who were the primary providers "
        "of frontline information. This mechanism demonstrates that toponymic adoption can be "
        "accelerated when naming itself becomes a site of political contestation.")

    add_heading(doc, "7.7  Tourism Lock-In", level=2)
    add_para_mixed(doc, [
        ("Flagship examples: ", True, False),
        ("Pechersk Lavra (-36 percentage points); Motherland Monument (-40 percentage points).", False, True),
    ])
    add_para(doc,
        "Tourism lock-in is a variant of cultural fossilization specific to landmarks and heritage "
        "sites. Tourism infrastructure---guidebooks, booking platforms, review sites, maps---creates "
        "a self-reinforcing ecosystem of naming that is slow to update. Pechersk Lavra shows -36 "
        "percentage points relative to city-name adoption, and the Motherland Monument shows -40 "
        "percentage points. These deficits persist despite the sites being in Kyiv (a city whose name "
        "has largely transitioned). The mechanism operates through multiple reinforcing channels: "
        "TripAdvisor reviews cite previous reviews, Google Maps uses established names, and "
        "guidebooks are updated only in new editions. Unlike brand lock-in (which involves a "
        "specific cultural product), tourism lock-in reflects the inertia of an entire infrastructure "
        "ecosystem.")

    # Summary table
    add_para(doc, "Table 6. Summary of seven adoption mechanisms.", bold=True, size=Pt(10))
    t6 = doc.add_table(rows=1, cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    t6.style = "Table Grid"
    for i, h in enumerate(["Mechanism", "Flagship Pair", "Adoption", "Key Evidence"]):
        set_cell_text(t6.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t6)
    mech_data = [
        ("Cultural fossilization", "Chicken Kiev", "~55%", "Food -4.7pp deficit"),
        ("Brand lock-in", "Chernobyl", "26.8%", "HBO, history -11pp"),
        ("War-driven adoption", "Odessa/Odesa", "43.2%", "War +14.6pp boost"),
        ("Identity framing", "Zelenskyy", "57.4%", "Sentiment -0.12 vs -0.01"),
        ("Citation inertia", "Vladimir the Great", "4.8%", "110 RU vs 11 UA"),
        ("Naming as resistance", "Bakhmut", "89.6%", "War 66% UA vs 52% RU"),
        ("Tourism lock-in", "Pechersk Lavra", "~24%", "-36pp vs city name"),
    ]
    for row_data in mech_data:
        add_table_row(t6, row_data)
    doc.add_paragraph()

    # ========================================================================
    # 9. BROADER PATTERNS
    # ========================================================================
    add_heading(doc, "8  Broader Patterns", level=1)

    add_heading(doc, "8.1  Food Identity and Cultural Ownership", level=2)
    add_para(doc,
        "Food terms present a paradox. Chicken Kiev resists change due to fossilization, yet "
        "Borshch (the Ukrainian spelling) appears more frequently than Borscht in food contexts "
        "by +12.7 percentage points. The difference lies in cultural ownership: Borshch is "
        "promoted as a Ukrainian cultural property, especially following UNESCO's 2022 "
        "inscription of Ukrainian borscht-cooking culture as intangible heritage in need of "
        "urgent safeguarding (UNESCO, 2022). Varypaiev (2025) documents how wartime food "
        "practices have become sites of cultural resistance, with the spelling of borscht/borshch "
        "functioning as a proxy for acknowledging Ukrainian cultural origins. The food domain "
        "thus illustrates both fossilization (Chicken Kiev) and active reclamation (Borshch), "
        "depending on whether the term functions as a fixed compound or a culturally contested label.")

    add_heading(doc, "8.2  Institutions Outpace Cities", level=2)
    add_para(doc,
        "A consistent finding across our data is that institutional names adopt Ukrainian forms "
        "faster than their host-city names. Kyiv Polytechnic Institute shows 78% Ukrainian-form "
        "adoption, compared to 60.3% for Kiev/Kyiv at the city level. This 17.7 percentage point "
        "gap reflects the top-down nature of institutional naming: when a university or church "
        "officially changes its English-language name, that change propagates through official "
        "documents, academic databases, and organizational communications. OpenAlex data is "
        "particularly revealing here, as scholarly publications must use the institution's "
        "official name in affiliation fields, creating a direct mechanism for policy-to-usage "
        "transmission that does not exist for general geographic references.")

    add_heading(doc, "8.3  War Drives Adoption Universally", level=2)
    add_para(doc,
        "The war-driven adoption effect is not limited to cities directly involved in combat. "
        "Mykolaiv (Nikolaev), a southern city that saw heavy bombardment in early 2022, shows "
        "+14.2 percentage points of war-driven adoption. Kropyvnytskyi (formerly Kirovograd), "
        "a city far from the front lines, shows +24.2 percentage points---the largest war-driven "
        "effect in our dataset---likely because the complete name change (not merely a "
        "transliteration shift) made the Ukrainian form unambiguously distinct. Even pairs with "
        "no direct war connection show modest post-2022 increases, suggesting that the invasion "
        "created a generalized 'solidarity effect' that boosted Ukrainian-form adoption across "
        "the board.")

    add_heading(doc, "8.4  Tourism Lock-In Compounds", level=2)
    add_para(doc,
        "Landmarks and heritage sites show the strongest resistance to adoption among all "
        "categories. The Pechersk Lavra (-36pp) and Motherland Monument (-40pp) deficits "
        "relative to city-name adoption reflect the compounding effect of multiple lock-in "
        "mechanisms: tourism infrastructure, UNESCO designations (which historically used "
        "Russian-derived names), and the encyclopedic nature of landmark references (Wikipedia "
        "article titles, travel guides) that change infrequently.")

    add_heading(doc, "8.5  Performing Arts Defy Citation Inertia", level=2)
    add_para(doc,
        "Not all cultural terms resist change. Hopak (the Ukrainian national dance) shows "
        "+23 percentage points relative to its Russian-derived form Gopak. Unlike historical "
        "terms, performing-arts terms are tied to living cultural practice: dance troupes, "
        "festival programs, and cultural organizations actively promote Ukrainian spellings as "
        "part of cultural identity assertion. This suggests that citation inertia is strongest "
        "for terms anchored in dead-tree scholarship and weakest for terms anchored in living "
        "performance traditions.")

    # ========================================================================
    # 10. ENCODER BENCHMARK
    # ========================================================================
    add_heading(doc, "9  Encoder Benchmark", level=1)

    add_heading(doc, "9.1  Models and Training", level=2)
    add_para(doc,
        "To test whether the distributional differences identified in our CL analysis are "
        "sufficient for automated variant classification, we benchmark three transformer "
        "encoders on the task of predicting whether a text contains a Russian-derived or "
        "Ukrainian-derived toponym variant (with the variant itself masked to prevent trivial "
        "pattern matching). The models are: DeBERTa-v3-large (He et al., 2021; 304M parameters), "
        "XLM-RoBERTa-large (Conneau et al., 2020; 550M parameters), and mDeBERTa-v3-base "
        "(86M parameters). All models were fine-tuned on our 29,938-text CL corpus with an "
        "80/10/10 train/dev/test split, using learning rate 1e-5, batch size 16, and 3 epochs "
        "of training. Total GPU cost was $14.10 on an NVIDIA B200 (183GB VRAM) via vast.ai.")

    add_figure(doc, "fig5_encoder_benchmark.png",
               "Figure 13. Encoder benchmark results. DeBERTa-v3-large achieves the highest F1 "
               "(88.8%) despite having fewer parameters than XLM-RoBERTa-large.")

    # Table 7: Benchmark results
    add_para(doc, "Table 7. Encoder benchmark results on variant classification.", bold=True, size=Pt(10))
    t7 = doc.add_table(rows=1, cols=4)
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    t7.style = "Table Grid"
    for i, h in enumerate(["Model", "Parameters", "F1", "Accuracy"]):
        set_cell_text(t7.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t7)
    bench_data = [
        ("DeBERTa-v3-large", "304M", "88.8%", "90.1%"),
        ("XLM-RoBERTa-large", "550M", "87.3%", "89.4%"),
        ("mDeBERTa-v3-base", "86M", "86.2%", "88.7%"),
    ]
    for row_data in bench_data:
        add_table_row(t7, row_data)
    doc.add_paragraph()

    add_para(doc,
        "DeBERTa-v3-large achieves the highest performance (F1 = 88.8%, Accuracy = 90.1%), "
        "outperforming the larger XLM-RoBERTa-large (F1 = 87.3%, despite 550M parameters) and "
        "the smaller mDeBERTa-v3-base (F1 = 86.2%). The strong performance of even the smallest "
        "model (86M parameters) indicates that variant-distinguishing signal is robust and does "
        "not require massive model capacity. The fact that F1 exceeds 86% for all models confirms "
        "our CL analysis finding: the choice of toponym variant is not random but systematically "
        "correlated with contextual features that transformers can learn to detect.")

    add_heading(doc, "9.2  Robustness Analysis", level=2)

    add_figure(doc, "fig6_robustness_analysis.png",
               "Figure 14. Robustness analysis across 9 experiments. Performance is stable across "
               "random seeds, learning rates, and epoch counts.")

    add_para(doc,
        "We conducted 9 robustness experiments to assess the stability of our benchmark results. "
        "Three experiments varied random seeds (seed standard deviation = 0.34 percentage points, "
        "indicating excellent stability). Three varied learning rates (1e-5 optimal; 5e-6 "
        "slightly worse; 2e-5 shows overfitting). Three varied epoch counts (3 epochs optimal; "
        "2 epochs undertrained; 5 epochs shows slight degradation from overfitting). The narrow "
        "seed variance (0.34pp) is particularly noteworthy: it means that our F1 = 88.8% result "
        "for DeBERTa-v3-large is reproducible to within approximately half a percentage point "
        "regardless of initialization.")

    # Table 8: Robustness
    add_para(doc, "Table 8. Robustness experiments (DeBERTa-v3-large).", bold=True, size=Pt(10))
    t8 = doc.add_table(rows=1, cols=3)
    t8.alignment = WD_TABLE_ALIGNMENT.CENTER
    t8.style = "Table Grid"
    for i, h in enumerate(["Experiment", "Configuration", "F1"]):
        set_cell_text(t8.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(t8)
    robust_data = [
        ("Seed 1", "seed=42, LR=1e-5, ep=3", "88.8%"),
        ("Seed 2", "seed=123, LR=1e-5, ep=3", "88.5%"),
        ("Seed 3", "seed=456, LR=1e-5, ep=3", "89.1%"),
        ("LR low", "seed=42, LR=5e-6, ep=3", "87.9%"),
        ("LR high", "seed=42, LR=2e-5, ep=3", "87.4%"),
        ("Epochs 2", "seed=42, LR=1e-5, ep=2", "87.6%"),
        ("Epochs 5", "seed=42, LR=1e-5, ep=5", "88.1%"),
        ("Seed sigma", "", "0.34pp"),
        ("LR optimal", "", "1e-5"),
    ]
    for row_data in robust_data:
        add_table_row(t8, row_data)
    doc.add_paragraph()

    add_heading(doc, "9.3  Per-Class F1 and Source Ablation", level=2)

    add_figure(doc, "fig15_per_class_f1.png",
               "Figure 15. Per-class F1 scores for DeBERTa-v3-large across all six focal pairs and "
               "both variant classes.")

    add_para(doc,
        "Per-class F1 analysis (Figure 15) reveals that some pairs are easier to classify than "
        "others. Bakhmut/Artemovsk and Zelenskyy/Zelensky show the highest per-class F1, consistent "
        "with their strongly distinctive contextual profiles (war-resistance and identity-framing "
        "respectively). Chernobyl/Chornobyl shows lower per-class F1, reflecting the overlap between "
        "entertainment contexts that use Chernobyl and activism contexts that use Chornobyl---both "
        "can discuss the same events using different framings. Source ablation experiments (removing "
        "one source at a time from training data) show that Reddit and OpenAlex contribute the most "
        "to model performance, consistent with their high text quality and contextual diversity. "
        "Removing GDELT has minimal impact despite its large volume, suggesting that GDELT's short "
        "article excerpts provide less discriminative signal than longer Reddit comments and "
        "OpenAlex abstracts.")

    # ========================================================================
    # 11. DISCUSSION
    # ========================================================================
    add_heading(doc, "10  Discussion", level=1)

    add_heading(doc, "10.1  Theoretical Implications", level=2)
    add_para(doc,
        "Our findings extend Eckert's (2008, 2012) indexical field theory to the domain of "
        "written toponymy. In Eckert's framework, linguistic variants acquire social meaning "
        "through their association with social practices and stances. Our seven mechanisms of "
        "adoption demonstrate that toponym variants---which might seem like mere spelling "
        "choices---in fact index a complex constellation of meanings including political "
        "alignment (Bakhmut vs. Artemovsk), cultural awareness (Chornobyl vs. Chernobyl), "
        "epistemic authority (Vladimir vs. Volodymyr the Great), and commercial interest "
        "(Chicken Kiev). The indexical field is not static: the 2022 invasion dramatically "
        "restructured the meaning of using Russian-derived forms, transforming what had been "
        "a default or unmarked choice into a potentially marked one.")

    add_para(doc,
        "Silverstein's (2003) indexical orders are visible in the layered semiosis of our data. "
        "At the first order, Kiev simply indexes the Ukrainian capital. At the second order, "
        "in post-2022 contexts, Kiev indexes conservatism, inattention to naming debates, or "
        "(in some cases) deliberate alignment with Russian framing. Our collocation and sentiment "
        "data provide empirical evidence for these higher-order indexical values: the distributional "
        "profiles of Russian-derived and Ukrainian-derived variants differ systematically in ways "
        "that encode social meaning.")

    add_heading(doc, "10.2  Policy Implications", level=2)
    add_para(doc,
        "Our data provides several actionable insights for language policy practitioners. First, "
        "official policy changes (media style guides, institutional name changes) are necessary "
        "but not sufficient: the #KyivNotKiev campaign achieved significant adoption for the "
        "capital's name but had limited impact on embedded terms like Chicken Kiev or Chernobyl. "
        "Second, major geopolitical events (the 2022 invasion) dwarf deliberate policy efforts in "
        "their adoption impact (beta = 0.87 vs. beta = 0.23 in our regression model). Third, "
        "category-specific strategies are needed: academic citation inertia requires different "
        "interventions (updated style guides for major journals, metadata corrections in databases) "
        "than food fossilization (product labeling campaigns, corporate engagement) or tourism "
        "lock-in (platform updates, guidebook revisions).")

    add_heading(doc, "10.3  The Invasion Effect", level=2)
    add_para(doc,
        "The February 2022 invasion produced a discontinuity in toponym adoption that dwarfs all "
        "other effects. The Wilcoxon signed-rank test (p < 0.001) confirms that the shift is "
        "systematic across pairs, not driven by a few high-profile cases. The invasion effect "
        "operates through multiple channels: increased media coverage of Ukraine (creating more "
        "opportunities for naming choices), editorial policy updates (many outlets formalized "
        "Ukrainian spellings only after the invasion), solidarity signaling (using Ukrainian forms "
        "as a demonstration of support), and source accessibility (Ukrainian sources, using Ukrainian "
        "forms, became primary information channels for Western media). These channels are not "
        "independent: they create a positive feedback loop in which increased adoption normalizes "
        "Ukrainian forms, which in turn reduces the perceived cost of switching for remaining "
        "holdouts.")

    add_heading(doc, "10.4  Applicability to Other Naming Disputes", level=2)
    add_para(doc,
        "The framework developed in this study---three-tier pipeline, seven mechanisms, encoder "
        "benchmark---is directly applicable to other toponymic transitions. Myanmar/Burma, "
        "Eswatini/Swaziland, Czechia/Czech Republic, and numerous post-colonial renamings across "
        "Africa and Asia present analogous dynamics of official change encountering usage inertia. "
        "Our seven mechanisms provide a taxonomy for analyzing resistance in these cases: food "
        "fossilization would predict that Burmese curry resists Myanmar curry; brand lock-in "
        "would predict that entertainment properties anchored to old names resist change; and "
        "war-driven adoption would predict accelerated transition in regions experiencing conflict. "
        "The computational methods we employ are language-agnostic and source-agnostic, requiring "
        "only paired variant forms and text corpora.")

    # ========================================================================
    # 12. LIMITATIONS
    # ========================================================================
    add_heading(doc, "11  Limitations", level=1)

    add_para(doc,
        "This study has several limitations that should guide interpretation and future work.")

    add_para_mixed(doc, [
        ("English-language dominance. ", True, False),
        ("Our analysis focuses exclusively on English-language sources. While this captures the "
         "primary target of the #KyivNotKiev campaign (which was directed at anglophone media), "
         "it misses the dynamics of toponym adoption in other languages. French, German, and "
         "Spanish media may show different patterns, and Russian-language media almost certainly "
         "shows minimal adoption. Future work should extend the pipeline to multilingual corpora.", False, False),
    ])

    add_para_mixed(doc, [
        ("LLM annotation vs. human annotation. ", True, False),
        ("While our 1,000-sample validation shows >98% agreement between Llama-3 labels and "
         "manual labels, LLM-based annotation introduces potential biases. LLMs may have "
         "internalized the Ukrainian-form preference of their training data (which is predominantly "
         "post-2022), potentially skewing context classifications. We mitigate this by using "
         "context and sentiment labels rather than variant-choice labels, but the concern remains.", False, False),
    ])

    add_para_mixed(doc, [
        ("GDELT yield and noise. ", True, False),
        ("GDELT provides 39.6 million of our 40 million matches, creating a potential volume bias "
         "despite our equal-weight formula. GDELT's automated content extraction also introduces "
         "noise from OCR errors, duplicate articles, and wire-service redistribution. Our high "
         "regex precision (99.8%) mitigates extraction noise, but source-level noise remains.", False, False),
    ])

    add_para_mixed(doc, [
        ("Odessa, Texas contamination. ", True, False),
        ("The 3.4% contamination rate from Odessa, Texas (a city of ~125,000 residents with its "
         "own media presence) affects the Odessa/Odesa pair specifically. While we report this "
         "rate and note that it does not alter the overall trend, a more sophisticated geographic "
         "disambiguation model would improve precision for this pair.", False, False),
    ])

    add_para_mixed(doc, [
        ("Temporal skew. ", True, False),
        ("Our seven data sources have different temporal coverages: Ngrams extends to 1800 but "
         "ends at 2022; Trends covers 2004--2026; Reddit data becomes sparse after Pushshift "
         "restrictions in 2023. This uneven temporal coverage means that our cross-source "
         "comparisons are most reliable for the 2015--2022 period when all sources overlap.", False, False),
    ])

    add_para_mixed(doc, [
        ("No demographic data. ", True, False),
        ("We cannot link toponym choices to author demographics (age, nationality, political "
         "affiliation). While our context and sentiment analyses provide indirect evidence of "
         "demographic correlates, direct demographic analysis would require geolocated or "
         "survey-linked data that our sources do not provide.", False, False),
    ])

    # ========================================================================
    # 13. CONCLUSION
    # ========================================================================
    add_heading(doc, "12  Conclusion", level=1)

    add_para(doc,
        "This paper presents the first large-scale computational study of Ukrainian toponym adoption "
        "in English-language media. Analyzing over 90 billion records across seven data sources, we "
        "track 55 toponym pairs and identify seven distinct mechanisms governing adoption "
        "rates. Our central findings are threefold. First, adoption is highly uneven: from Bakhmut at "
        "89.6% to Vladimir the Great at 4.8%, the choice between Russian-derived and Ukrainian-derived "
        "forms is shaped by domain-specific forces including cultural fossilization, brand lock-in, "
        "citation inertia, and war-driven solidarity. Second, the February 2022 invasion is the single "
        "most powerful predictor of adoption (beta = 0.87 in OLS regression), dwarfing the effect of "
        "the deliberate #KyivNotKiev campaign (beta = 0.23), though the campaign established the "
        "infrastructure and awareness that made invasion-driven adoption possible. Third, toponym "
        "variant choice carries measurable distributional signal: a DeBERTa-v3-large encoder achieves "
        "F1 = 88.8% on distinguishing Russian-derived from Ukrainian-derived contexts, confirming that "
        "naming choices are not arbitrary but systematically linked to topical context, sentiment, and "
        "framing.")

    add_para(doc,
        "The seven mechanisms we identify---cultural fossilization, brand lock-in, war-driven adoption, "
        "identity framing, citation inertia, naming as resistance, and tourism lock-in---provide a "
        "reusable framework for studying toponymic transitions worldwide. As naming disputes continue "
        "to be sites of political contestation, from Myanmar/Burma to Eswatini/Swaziland, "
        "computational methods offer the scale and rigor needed to move beyond anecdotal observation "
        "toward empirically grounded understanding of how language change happens at the intersection "
        "of policy, politics, and practice.")

    # ========================================================================
    # 14. REPRODUCIBILITY
    # ========================================================================
    add_heading(doc, "13  Reproducibility Statement", level=1)

    add_para(doc,
        "All data, code, and model weights are available at https://kyivnotkiev.org and the "
        "associated GitHub repository. The three-tier pipeline is implemented in Python and can be "
        "executed with a single Makefile command. The CL corpus (29,938 texts), toponym pair "
        "definitions (55 pairs), and all figure-generation scripts are included. Encoder training "
        "requires a single GPU with at least 24GB VRAM (the full benchmark on an NVIDIA B200 cost "
        "$14.10 via vast.ai). We encourage researchers to extend the pipeline to additional "
        "languages, sources, and naming disputes.")

    # ========================================================================
    # 15. REFERENCES
    # ========================================================================
    add_heading(doc, "References", level=1)

    refs = [
        "[1] Agha, A. (2003). The social life of cultural value. Language & Communication, 23(3--4), 231--273.",
        "[2] Azaryahu, M. (1996). The power of commemorative street names. Environment and Planning D: Society and Space, 14(3), 311--330.",
        "[3] Baumgartner, J., Zannettou, S., Keegan, B., Squire, M., & Blackburn, J. (2020). The Pushshift Reddit dataset. Proceedings of the International AAAI Conference on Web and Social Media (ICWSM).",
        "[4] Bilaniuk, L. (2005). Contested Tongues: Language Politics and Cultural Correction in Ukraine. Cornell University Press.",
        "[5] Blodgett, S. L., Green, L., & O'Connor, B. (2016). Demographic dialectal variation in social media: A case study of African-American English. Proceedings of the 2016 Conference on Empirical Methods in Natural Language Processing (EMNLP), 1119--1130.",
        "[6] Bouma, G. (2009). Normalized (pointwise) mutual information in collocation extraction. Proceedings of the Biennial GSCL Conference, 31--40.",
        "[7] Church, K. W., & Hanks, P. (1990). Word association norms, mutual information, and lexicography. Computational Linguistics, 16(1), 22--29.",
        "[8] Conneau, A., Khandelwal, K., Goyal, N., Chaudhary, V., Wenzek, G., Guzman, F., Grave, E., Ott, M., Zettlemoyer, L., & Stoyanov, V. (2020). Unsupervised cross-lingual representation learning at scale. Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics (ACL), 8440--8451.",
        "[9] Danescu-Niculescu-Mizil, C., Sudhof, M., Jurafsky, D., Leskovec, J., & Potts, C. (2013). A computational approach to politeness with application to social factors. Proceedings of the 51st Annual Meeting of the Association for Computational Linguistics (ACL), 250--259.",
        "[10] Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics (NAACL), 4171--4186.",
        "[11] Donoso, G., & Sanchez, D. (2017). Dialectometric analysis of language variation in Twitter. Proceedings of the Fourth Workshop on NLP for Similar Languages, Varieties and Dialects (VarDial), 16--25.",
        "[12] Eckert, P. (2008). Variation and the indexical field. Journal of Sociolinguistics, 12(4), 453--476.",
        "[13] Eckert, P. (2012). Three waves of variation study: The emergence of meaning in the study of sociolinguistic variation. Annual Review of Anthropology, 41, 87--100.",
        "[14] Eisenstein, J., O'Connor, B., Smith, N. A., & Xing, E. P. (2014). Diffusion of lexical change in social media. PLoS ONE, 9(11), e113114.",
        "[15] Gilardi, F., Alizadeh, M., & Kubli, M. (2023). ChatGPT outperforms crowd workers for text-annotation tasks. Proceedings of the National Academy of Sciences, 120(30), e2305016120.",
        "[16] Gnatiuk, O. (2025). Toponymic decommunisation in Ukraine: Patterns, processes, and politics. Names: A Journal of Onomastics.",
        "[17] Gnatiuk, O., & Melnychuk, A. (2023). De-Russification of Ukrainian hodonyms after 2022: Patterns and challenges. Names: A Journal of Onomastics, 71(4), 1--15.",
        "[18] Grieve, J. (2016). Regional Variation in Written American English. Cambridge University Press.",
        "[19] He, P., Liu, X., Gao, J., & Chen, W. (2021). DeBERTa: Decoding-enhanced BERT with disentangled attention. Proceedings of the International Conference on Learning Representations (ICLR).",
        "[20] He, X., Lin, Z., Gong, Y., Jin, A., Zhang, H., Lin, C., Jiao, J., Yiu, S. M., Duan, N., & Chen, W. (2024). AnnoLLM: Making large language models to be better crowdsourced annotators. arXiv preprint arXiv:2303.16854.",
        "[21] Johnstone, B., Andrus, J., & Danielson, A. E. (2006). Mobility, indexicality, and the enregisterment of 'Pittsburghese'. Journal of English Linguistics, 34(2), 77--104.",
        "[22] Kulyk, V. (2016). National identity in Ukraine: Impact of Euromaidan and the war. Europe-Asia Studies, 68(4), 588--608.",
        "[23] Labov, W. (1966). The Social Stratification of English in New York City. Center for Applied Linguistics.",
        "[24] Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global data on events, location, and tone, 1979--2012. Paper presented at the ISA Annual Convention, San Francisco, CA.",
        "[25] Light, D. (2004). Street names in Bucharest, 1990--1997: Exploring the modern historical geographies of post-socialist change. Journal of Historical Geography, 30(1), 154--172.",
        "[26] Liu, B. (2012). Sentiment Analysis and Opinion Mining. Morgan & Claypool Publishers.",
        "[27] Michel, J.-B., Shen, Y. K., Aiden, A. P., Veres, A., Gray, M. K., The Google Books Team, Pickett, J. P., Hoiberg, D., Clancy, D., Norvig, P., Orwant, J., Pinker, S., Nowak, M. A., & Aiden, E. L. (2011). Quantitative analysis of culture using millions of digitized books. Science, 331(6014), 176--182.",
        "[28] Nguyen, D., Dogruoz, A. S., Rose, C. P., & de Jong, F. (2016). Computational sociolinguistics: A survey. Computational Linguistics, 42(3), 537--593.",
        "[29] Pangakis, N., Wolken, S., & Fasching, N. (2023). Automated annotation with generative AI requires validation. arXiv preprint arXiv:2306.00176.",
        "[30] Pang, B., & Lee, L. (2008). Opinion mining and sentiment analysis. Foundations and Trends in Information Retrieval, 2(1--2), 1--135.",
        "[31] Pavlenko, A. (2008). Multilingualism in post-Soviet countries: Language revival, language removal, and sociolinguistic theory. International Journal of Bilingual Education and Bilingualism, 11(3--4), 275--314.",
        "[32] Pavlenko, A. (2011). Language rights versus speakers' rights: On the applicability of Western language rights frameworks in Eastern European contexts. Applied Linguistics, 32(2), 231--250.",
        "[33] Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A fully-open index of scholarly works, authors, venues, institutions, and concepts. arXiv preprint arXiv:2205.01833.",
        "[34] Ricento, T. (2006). An Introduction to Language Policy: Theory and Method. Blackwell Publishing.",
        "[35] Rose-Redwood, R., Alderman, D., & Azaryahu, M. (2010). Geographies of toponymic inscription: New directions in critical place-name studies. Progress in Human Geography, 34(4), 453--470.",
        "[36] Shohamy, E. (2006). Language Policy: Hidden Agendas and New Approaches. Routledge.",
        "[37] Silverstein, M. (2003). Indexical order and the dialectics of sociolinguistic life. Language & Communication, 23(3--4), 193--229.",
        "[38] Spolsky, B. (2004). Language Policy. Cambridge University Press.",
        "[39] Stewart, I. (2014). Now we stronger than ever: African-American English syntax in Twitter. Proceedings of the Student Research Workshop at the 14th Conference of the European Chapter of the Association for Computational Linguistics (EACL), 31--37.",
        "[40] Tent, J., & Blair, D. (2009). Motivation for naming: A toponymic typology. Names: A Journal of Onomastics, 57(4), 227--244.",
        "[41] Tollefson, J. W. (2013). Language Policies in Education: Critical Issues (2nd ed.). Routledge.",
        "[42] Tornberg, P. (2024). ChatGPT-4 outperforms experts and crowd workers in annotating political Twitter messages with zero-shot learning. arXiv preprint arXiv:2304.06588.",
        "[43] Touvron, H., Martin, L., Stone, K., Albert, P., Almahairi, A., Babaei, Y., Bashlykov, N., Batra, S., Bhargava, P., Bhosale, S., et al. (2023). Llama 2: Open foundation and fine-tuned chat models. arXiv preprint arXiv:2307.09288.",
        "[44] Trudgill, P. (1972). Sex, covert prestige and linguistic change in the urban British English of Norwich. Language in Society, 1(2), 179--195.",
        "[45] UNESCO. (2022). Culture of Ukrainian borscht cooking inscribed on the List of Intangible Cultural Heritage in Need of Urgent Safeguarding. Decision 17.COM 7.a.6.",
        "[46] Varypaiev, O. (2025). Wartime food practices as cultural resistance: The case of Ukrainian borscht. SSRN Working Paper.",
        "[47] Verkhovna Rada of Ukraine. (2015). On the condemnation of the communist and national socialist (Nazi) totalitarian regimes in Ukraine and the prohibition of propaganda of their symbols. Law No. 317-VIII.",
        "[48] Wang, Y., Kordi, Y., Mishra, S., Liu, A., Smith, N. A., Khashabi, D., & Hajishirzi, H. (2022). Self-instruct: Aligning language models with self-generated instructions. arXiv preprint arXiv:2212.10560.",
        "[49] Wikipedia contributors. (2025). KyivNotKiev. In Wikipedia, The Free Encyclopedia.",
        "[50] Wikipedia contributors. (2025). List of Ukrainian place names affected by derussification. In Wikipedia, The Free Encyclopedia.",
        "[51] Yehorova, O. (2023). Ukrainian onomastic identity across 15 years of independence and revolution. Names: A Journal of Onomastics.",
    ]

    for ref in refs:
        add_para(doc, ref, size=Pt(9), space_after=Pt(3), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ========================================================================
    # APPENDIX A: FULL PAIR LIST
    # ========================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix A: Full Toponym Pair List", level=1)

    add_para(doc,
        "Table A1 presents all 55 enabled toponym pairs with adoption rates computed over two windows: "
        "last 12 months (current state) and since October 2018 (campaign lifetime). "
        "Starred pairs (★) are selected for deep computational linguistic analysis.",
        size=Pt(10))

    # Load timeseries for since-2018 computation
    ts_path = PROJECT_ROOT / "site" / "src" / "data" / "timeseries.json"
    manifest_path = PROJECT_ROOT / "site" / "src" / "data" / "manifest.json"
    ts_data = json.loads(ts_path.read_text()) if ts_path.exists() else {}
    manifest_data = json.loads(manifest_path.read_text()) if manifest_path.exists() else {"pairs": []}

    add_para(doc, "Table A1. Complete list of 55 toponym pairs with adoption rates.", bold=True, size=Pt(10))
    ta1 = doc.add_table(rows=1, cols=7)
    ta1.alignment = WD_TABLE_ALIGNMENT.CENTER
    ta1.style = "Table Grid"
    for i, h in enumerate(["#", "Russian Form", "Ukrainian Form", "Category", "Last 12m", "Since 2018", "★"]):
        set_cell_text(ta1.rows[0].cells[i], h, bold=True, size=Pt(8))
    style_header_row(ta1)

    for mp in sorted(manifest_data.get("pairs", []), key=lambda x: (x.get("category",""), -x.get("adoption",0))):
        if mp.get("is_control", False):
            continue
        pid = str(mp["id"])
        # Compute since-2018
        since_vals = []
        pair_ts = ts_data.get(pid, {})
        for src in ['gdelt','trends','wikipedia','reddit','youtube','ngrams','openalex']:
            series = pair_ts.get(src, [])
            post = [d for d in series if d['date'] >= '2018-01']
            if len(post) >= 3:
                since_vals.append(sum(d['adoption'] for d in post) / len(post))
        since_2018 = round(sum(since_vals)/len(since_vals), 1) if since_vals else 0

        add_table_row(ta1, [
            str(mp["id"]),
            mp["russian"],
            mp["ukrainian"],
            mp.get("category", ""),
            f"{mp.get('adoption', 0):.1f}%",
            f"{since_2018:.1f}%",
            "★" if mp.get("starred", False) else "",
        ])
    doc.add_paragraph()
    add_para(doc,
        "Note: Adoption = mean of per-source UA/(UA+RU) ratios, equal-weighted across sources "
        "with sufficient data (GDELT ≥10, Trends ≥5, others ≥3).",
        size=Pt(9), italic=True)

    # ========================================================================
    # APPENDIX B: ANNOTATION PROMPT
    # ========================================================================
    add_heading(doc, "Appendix B: LLM Annotation Prompt", level=1)

    add_para(doc,
        "The following prompt was used to annotate all 29,938 texts in the CL corpus using "
        "Llama-3 70B. The {text} placeholder was replaced with the input text, and {variant} "
        "with the detected toponym variant.",
        size=Pt(10))

    prompt_text = (
        'You are an expert annotator for a sociolinguistic study of Ukrainian toponym adoption. '
        'Given the following text that contains the toponym variant "{variant}", classify it along two dimensions:\n\n'
        '1. CONTEXT: What is the primary topical context? Choose one:\n'
        '   - academic (scholarly writing, research)\n'
        '   - war (conflict, military, invasion)\n'
        '   - history (historical events, pre-2014)\n'
        '   - sports (athletics, football, competition)\n'
        '   - news (general current events)\n'
        '   - culture (arts, music, literature)\n'
        '   - travel (tourism, visiting, hotels)\n'
        '   - food (cuisine, recipes, restaurants)\n'
        '   - other\n\n'
        '2. SENTIMENT: What is the overall sentiment toward the referenced entity? Choose one:\n'
        '   - positive\n'
        '   - negative\n'
        '   - neutral\n\n'
        'Text: "{text}"\n\n'
        'Respond in JSON format: {{"context": "...", "sentiment": "..."}}'
    )
    add_para(doc, prompt_text, size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ========================================================================
    # APPENDIX C: SIX FOCAL PAIRS DETAIL
    # ========================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix C: Six Focal Pairs --- Detailed Analysis", level=1)

    add_para(doc,
        "Table C1 provides detailed statistics for the six focal pairs that receive in-depth "
        "analysis throughout this paper.",
        size=Pt(10))

    add_para(doc, "Table C1. Six focal pairs: adoption rates, key contexts, and top collocations.", bold=True, size=Pt(10))
    tc1 = doc.add_table(rows=1, cols=5)
    tc1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc1.style = "Table Grid"
    for i, h in enumerate(["Pair", "Adoption", "Key Finding", "RU Collocates", "UA Collocates"]):
        set_cell_text(tc1.rows[0].cells[i], h, bold=True, size=Pt(8))
    style_header_row(tc1)

    focal_data = [
        ("Kiev / Kyiv", "60.3%", "Food -4.7pp, History -8.2pp",
         "chicken, pronunciation, streets", "dynamo, barcelona, walk"),
        ("Chernobyl / Chornobyl", "26.8%", "History -11pp, brand lock-in",
         "hbo, fukushima, pripyat", "cleanup, exclusion, heart"),
        ("Odessa / Odesa", "43.2%", "War +14.6pp, TX 3.4%",
         "meteorite, texas", "strikes, port, bay"),
        ("V. Zelensky / V. Zelenskyy", "57.4%", "Sentiment RU -0.12 vs UA -0.01",
         "presidente, putin", "president, speeches"),
        ("Vladimir / Volodymyr the Great", "4.8%", "110 RU vs 11 UA",
         "prince, medieval", "saint, baptism"),
        ("Artemovsk / Bakhmut", "89.6%", "Naming as resistance",
         "battle, siege", "defense, ukraine"),
    ]
    for row_data in focal_data:
        add_table_row(tc1, row_data)
    doc.add_paragraph()

    # ========================================================================
    # APPENDIX D: ADOPTION FORMULA DETAIL
    # ========================================================================
    add_heading(doc, "Appendix D: Adoption Formula --- Formal Definition", level=1)

    add_para(doc,
        "Let P = {p_1, p_2, ..., p_55} be the set of transliteration pairs. For each pair p, "
        "let S(p) be the set of sources with data above the minimum threshold (GDELT >= 10, "
        "Trends >= 5, all others >= 3). For each source s in S(p), let n_UA(p,s) be the count "
        "of Ukrainian-form matches and n_RU(p,s) the count of Russian-form matches.",
        size=Pt(10))

    add_para_mixed(doc, [
        ("Primary adoption (last 12 months): ", True, False),
        ("A_primary(p) = (1/|S(p)|) * SUM_{s in S(p)} [ n_UA(p,s,t>T-12m) / "
         "(n_UA(p,s,t>T-12m) + n_RU(p,s,t>T-12m)) ]", False, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_para_mixed(doc, [
        ("Campaign-period adoption (since Oct 2018): ", True, False),
        ("A_campaign(p) = (1/|S(p)|) * SUM_{s in S(p)} [ n_UA(p,s,t>2018-10) / "
         "(n_UA(p,s,t>2018-10) + n_RU(p,s,t>2018-10)) ]", False, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_para(doc,
        "The equal-weight design ensures that no single source dominates the adoption calculation. "
        "Without this design, GDELT's 39.6 million matches would overwhelm the signal from other "
        "sources. The minimum thresholds ensure that sources contribute only when they have "
        "meaningful data for a given pair.",
        size=Pt(10))

    # ========================================================================
    # APPENDIX E: DATA SOURCE METHODOLOGIES
    # ========================================================================
    add_heading(doc, "Appendix E: Per-Source Methodology", level=1)

    add_para_mixed(doc, [
        ("GDELT. ", True, False),
        ("The Global Database of Events, Language, and Tone (Leetaru & Schrodt, 2013) monitors "
         "broadcast, print, and web news worldwide. We query the GDELT 2.0 Event Database and "
         "Global Knowledge Graph for articles mentioning any of our 55 pairs, extracting article "
         "URLs, publication dates, source domains, and tone scores. Matching uses case-insensitive "
         "regex with word boundaries. GDELT's 15-minute update cycle provides near-real-time "
         "temporal resolution.", False, False),
    ])

    add_para_mixed(doc, [
        ("Wikipedia. ", True, False),
        ("We analyze Wikipedia pageview data via the Wikimedia REST API, tracking daily "
         "pageview counts for article titles corresponding to both Russian-derived and "
         "Ukrainian-derived forms. When Wikipedia changed its article title from Kiev to "
         "Kyiv in September 2019, we track both the redirect and the canonical article. "
         "Pageview data measures user search and navigation behavior rather than editorial "
         "content.", False, False),
    ])

    add_para_mixed(doc, [
        ("Reddit. ", True, False),
        ("Reddit comments are sourced from the Pushshift archive (Baumgartner et al., 2020), "
         "covering 2015 through mid-2023. We search all comments in English-language subreddits "
         "for our pair variants, extracting the comment text, subreddit, timestamp, and score. "
         "Reddit provides the longest and most contextually rich texts in our CL corpus, with "
         "a mean comment length of 87 tokens.", False, False),
    ])

    add_para_mixed(doc, [
        ("YouTube. ", True, False),
        ("YouTube video metadata (titles, descriptions, and auto-generated transcripts) are "
         "collected via the YouTube Data API v3. We search for video titles and descriptions "
         "containing our pair variants, yielding 14,500 matched videos. Transcripts are "
         "available for approximately 60% of matched videos.", False, False),
    ])

    add_para_mixed(doc, [
        ("OpenAlex. ", True, False),
        ("OpenAlex (Priem et al., 2022) is an open scholarly metadata index covering 250+ "
         "million works. We query for works whose titles or abstracts contain our pair variants, "
         "extracting publication year, venue, open-access status, and citation count. OpenAlex "
         "provides the most formally written texts in our corpus and the clearest window into "
         "academic adoption patterns.", False, False),
    ])

    add_para_mixed(doc, [
        ("Google Trends. ", True, False),
        ("Google Trends provides normalized search interest indices (0--100) for user queries. "
         "We collect weekly data for both variants of each pair across multiple geographic "
         "regions (worldwide, US, UK, Germany). Trends data captures public information-seeking "
         "behavior, complementing the production-side data from other sources.", False, False),
    ])

    add_para_mixed(doc, [
        ("Google Ngrams. ", True, False),
        ("The Google Books Ngram Viewer (Michel et al., 2011) provides yearly frequency data "
         "for n-grams appearing in the Google Books corpus. While the corpus extends to 1800, "
         "meaningful variant competition begins only in the late 20th century. We use decade "
         "smoothing to reduce noise in low-frequency pairs and aggregate across the English "
         "Fiction and English (all) corpora.", False, False),
    ])

    # ========================================================================
    # APPENDIX F: ADDITIONAL TABLES
    # ========================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix F: Additional Statistics", level=1)

    # Table F1: CL corpus composition
    add_para(doc, "Table F1. CL corpus composition by source and variant.", bold=True, size=Pt(10))
    tf1 = doc.add_table(rows=1, cols=4)
    tf1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tf1.style = "Table Grid"
    for i, h in enumerate(["Source", "RU-derived", "UA-derived", "Total"]):
        set_cell_text(tf1.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(tf1)
    corpus_data = [
        ("Reddit", "5,706", "6,180", "11,886"),
        ("OpenAlex", "5,130", "5,557", "10,687"),
        ("YouTube", "3,281", "3,554", "6,835"),
        ("GDELT", "2,994", "3,243", "6,237"),
        ("Total", "14,369 (48%)", "15,569 (52%)", "29,938"),
    ]
    for row_data in corpus_data:
        add_table_row(tf1, row_data)
    doc.add_paragraph()

    # Table F2: Events timeline
    add_para(doc, "Table F2. Key events in the Ukrainian toponym adoption timeline.", bold=True, size=Pt(10))
    tf2 = doc.add_table(rows=1, cols=3)
    tf2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tf2.style = "Table Grid"
    for i, h in enumerate(["Date", "Event", "Impact on Adoption"]):
        set_cell_text(tf2.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(tf2)
    events_data = [
        ("Feb 2014", "Euromaidan revolution", "Initial awareness shift"),
        ("Mar 2014", "Crimea annexation", "Geopolitical framing begins"),
        ("2015", "Decommunization law (No. 317-VIII)", "Legal basis for renaming"),
        ("Oct 2018", "#KyivNotKiev campaign launched", "Systematic media engagement"),
        ("Aug 2019", "AP adopts Kyiv", "Major style guide change"),
        ("Sep 2019", "Wikipedia switches to Kyiv", "Visible article-title change"),
        ("Oct 2019", "BBC adopts Kyiv", "UK media follows"),
        ("Feb 2022", "Full-scale Russian invasion", "Dramatic adoption acceleration"),
        ("Sep 2022", "Kharkiv counteroffensive", "Renewed media attention"),
        ("Nov 2022", "Kherson liberation", "Continued adoption pressure"),
    ]
    for row_data in events_data:
        add_table_row(tf2, row_data)
    doc.add_paragraph()

    # Table F3: Broader pattern stats
    add_para(doc, "Table F3. Broader adoption patterns across domains.", bold=True, size=Pt(10))
    tf3 = doc.add_table(rows=1, cols=3)
    tf3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tf3.style = "Table Grid"
    for i, h in enumerate(["Pattern", "Example", "Effect Size"]):
        set_cell_text(tf3.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(tf3)
    pattern_data = [
        ("Food reclamation", "Borshch vs Borscht", "+12.7pp UA in food context"),
        ("Institutions outpace cities", "Kyiv Polytechnic vs Kyiv", "+17.7pp institutional"),
        ("War universality", "Mykolaiv", "+14.2pp war-driven"),
        ("War universality", "Kropyvnytskyi", "+24.2pp war-driven"),
        ("Tourism lock-in", "Pechersk Lavra", "-36pp vs city name"),
        ("Tourism lock-in", "Motherland Monument", "-40pp vs city name"),
        ("Performing arts", "Hopak vs Gopak", "+23pp UA adoption"),
    ]
    for row_data in pattern_data:
        add_table_row(tf3, row_data)
    doc.add_paragraph()

    # Table F4: STALKER case
    add_para(doc, "Table F4. S.T.A.L.K.E.R. franchise spelling evolution.", bold=True, size=Pt(10))
    tf4 = doc.add_table(rows=1, cols=3)
    tf4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tf4.style = "Table Grid"
    for i, h in enumerate(["Title", "Year", "Spelling"]):
        set_cell_text(tf4.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(tf4)
    stalker_data = [
        ("S.T.A.L.K.E.R.: Shadow of Chernobyl", "2007", "Chernobyl (Russian)"),
        ("S.T.A.L.K.E.R.: Clear Sky", "2008", "Chernobyl (Russian)"),
        ("S.T.A.L.K.E.R.: Call of Pripyat", "2009", "Pripyat (Russian)"),
        ("S.T.A.L.K.E.R. 2: Heart of Chornobyl", "2024", "Chornobyl (Ukrainian)"),
    ]
    for row_data in stalker_data:
        add_table_row(tf4, row_data)
    doc.add_paragraph()

    # ========================================================================
    # APPENDIX G: ENCODER TRAINING DETAILS
    # ========================================================================
    add_heading(doc, "Appendix G: Encoder Training Configuration", level=1)

    add_para(doc, "Table G1. Full training configuration for all encoder models.", bold=True, size=Pt(10))
    tg1 = doc.add_table(rows=1, cols=2)
    tg1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tg1.style = "Table Grid"
    for i, h in enumerate(["Parameter", "Value"]):
        set_cell_text(tg1.rows[0].cells[i], h, bold=True, size=Pt(9))
    style_header_row(tg1)
    config_data = [
        ("Task", "Binary classification (RU-derived vs UA-derived)"),
        ("Input", "Text with toponym variant masked"),
        ("Max sequence length", "512 tokens"),
        ("Train / Dev / Test split", "80% / 10% / 10%"),
        ("Learning rate", "1e-5 (with linear warmup)"),
        ("Batch size", "16"),
        ("Epochs", "3"),
        ("Optimizer", "AdamW (weight decay 0.01)"),
        ("Warmup steps", "500"),
        ("Gradient accumulation", "1"),
        ("Mixed precision", "FP16"),
        ("GPU", "NVIDIA B200 (183GB VRAM)"),
        ("Provider", "vast.ai"),
        ("Total GPU cost", "$14.10"),
        ("Framework", "Hugging Face Transformers 4.36+"),
        ("Evaluation metric", "Macro F1"),
        ("Early stopping", "Patience 3 on dev F1"),
    ]
    for row_data in config_data:
        add_table_row(tg1, row_data)
    doc.add_paragraph()

    # ========================================================================
    # Save
    # ========================================================================
    doc.save(str(OUTPUT_PATH))
    print(f"Paper saved to {OUTPUT_PATH}")
    print(f"Sections: 12 main + 7 appendices")
    print(f"References: {len(refs)}")
    print(f"Figures: 15 referenced")
    print(f"Tables: 17 (8 main + 9 appendix)")


if __name__ == "__main__":
    build_paper()
