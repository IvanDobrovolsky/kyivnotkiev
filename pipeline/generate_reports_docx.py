#!/usr/bin/env python3
"""Generate a single DOCX file containing all 59 toponym pair audit reports.

Reads from site/src/data/*.json, reuses make_svg_chart from generate_reports.py,
converts SVG charts to PNG via cairosvg, and embeds them in a professional DOCX.

Usage:
    python3 pipeline/generate_reports_docx.py
"""

import ctypes.util
import io
import json
import os
import re
import sys
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

# ── Ensure Homebrew cairo is findable (macOS) ──────────────────────────
_brew_lib = "/opt/homebrew/lib"
if os.path.isdir(_brew_lib):
    os.environ.setdefault("DYLD_LIBRARY_PATH", _brew_lib)
    # Also patch ctypes.util so cairocffi can find the dylib
    _orig_find = ctypes.util.find_library
    def _patched_find(name):
        result = _orig_find(name)
        if result is None and name == "cairo":
            candidate = os.path.join(_brew_lib, "libcairo.2.dylib")
            if os.path.exists(candidate):
                return candidate
        return result
    ctypes.util.find_library = _patched_find

# ── Project paths ──────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "site" / "src" / "data"
OUT  = ROOT / "output"
OUT.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT / "pair_reports.docx"

# ── Add pipeline dir to path so we can import from generate_reports ────
sys.path.insert(0, str(ROOT / "pipeline"))
from generate_reports import (
    make_svg_chart, fmt_number,
    SOURCE_NAMES, SOURCE_COLORS, EVENTS, FAMILY_COLORS,
)

# ── Third-party imports ────────────────────────────────────────────────
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Constants ──────────────────────────────────────────────────────────
BLUE_UA  = RGBColor(0x00, 0x57, 0xB8)
RED_RU   = RGBColor(0xD5, 0x2B, 0x1E)
GRAY     = RGBColor(0x6B, 0x72, 0x80)
PURPLE   = RGBColor(0x7C, 0x3A, 0xED)
GREEN    = RGBColor(0x05, 0x96, 0x69)
DARK     = RGBColor(0x1F, 0x29, 0x37)
LIGHT_BG = RGBColor(0xF4, 0xF5, 0xF7)

SOURCE_ORDER = ["trends", "gdelt", "wikipedia", "reddit",
                "youtube", "ngrams", "openalex", "openlibrary"]

FAM_ORDER = ["Anthropic Claude", "OpenAI GPT", "Google Gemini", "Google Gemma",
             "xAI Grok", "Meta Llama", "Mistral", "Alibaba Qwen"]

CTX_LABELS = {
    "war_conflict": "War/Conflict", "academic_science": "Academic",
    "politics": "Politics", "history": "History",
    "travel_tourism": "Travel/Tourism", "sports": "Sports",
    "general_news": "News", "culture_arts": "Culture/Arts",
    "food_cuisine": "Food/Cuisine", "business_economy": "Business",
    "religion": "Religion",
}

DENOM_SHORT = {
    "Moscow Patriarchate, Department for External Church Relations": "Moscow Patriarchate",
    "World Council of Churches": "WCC",
    "Ecumenical Patriarchate of Constantinople": "Constantinople",
    "Holy See / Vatican": "Vatican",
    "Russian Orthodox Church (patriarchia.ru)": "ROC (patriarchia.ru)",
}


# ── Helpers ────────────────────────────────────────────────────────────

def load_json(name):
    with open(DATA / name) as f:
        return json.load(f)


def strip_html(text):
    """Remove HTML tags from text, keeping content."""
    text = re.sub(r'<a\s[^>]*href=["\']([^"\']*)["\'][^>]*>(.*?)</a>', r'\2 (\1)', text)
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&times;', 'x', text)
    text = re.sub(r'&#\d+;', '', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def set_cell_shading(cell, hex_color):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="D1D5DB", width=4):
    """Set thin borders on entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{width}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def make_header_row(table, texts, bold=True, bg="E5E7EB"):
    """Format the first row of a table as a header."""
    row = table.rows[0]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = bold
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        set_cell_shading(cell, bg)


def add_data_row(table, texts, alignments=None, bold_cols=None):
    """Add a data row to a table."""
    row = table.add_row()
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        if bold_cols and i in bold_cols:
            run.bold = True
        if alignments and i < len(alignments):
            p.alignment = alignments[i]
    return row


def svg_to_png_bytes(svg_str, width=1800):
    """Convert an SVG string to PNG bytes using cairosvg."""
    # Wrap bare SVG if needed
    if not svg_str.strip().startswith('<?xml') and not svg_str.strip().startswith('<svg'):
        return None
    # Ensure the SVG has xmlns
    if 'xmlns=' not in svg_str:
        svg_str = svg_str.replace('<svg ', '<svg xmlns="http://www.w3.org/2000/svg" ', 1)
    # Extract viewBox and add explicit width/height (cairosvg needs them)
    vb = re.search(r'viewBox="([^"]+)"', svg_str)
    if vb:
        parts = vb.group(1).split()
        if len(parts) == 4:
            vb_w, vb_h = parts[2], parts[3]
            svg_str = svg_str.replace('viewBox=', f'width="{vb_w}" height="{vb_h}" viewBox=')
    # Remove style attr on root svg (display:block;width:100% confuses cairosvg)
    svg_str = re.sub(r'(<svg[^>]*?) style="[^"]*"', r'\1', svg_str)
    try:
        png_bytes = cairosvg.svg2png(
            bytestring=svg_str.encode('utf-8'),
            output_width=width,
            background_color="white",
        )
        return png_bytes
    except Exception as e:
        print(f"    SVG→PNG conversion error: {e}")
        return None


def convert_chart_batch(chart_jobs):
    """Convert a batch of SVG charts to PNG in parallel (max 3 workers)."""
    results = {}
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = {}
        for key, svg_str in chart_jobs.items():
            futures[executor.submit(svg_to_png_bytes, svg_str)] = key
        for future in as_completed(futures):
            key = futures[future]
            try:
                results[key] = future.result()
            except Exception as e:
                print(f"    Chart conversion error for {key}: {e}")
                results[key] = None
    return results


# ── Main document generation ───────────────────────────────────────────

def build_docx():
    print("Loading data files...")
    manifest     = load_json("manifest.json")
    timeseries   = load_json("timeseries.json")
    llm_raw      = load_json("llm_per_pair.json")
    llm_per_pair = llm_raw.get("pairs", {})
    collocations = load_json("cl_collocations.json")
    religious    = load_json("religious.json")
    pair_events  = load_json("pair_events.json")
    stats_raw    = load_json("statistical_tests.json")
    stats        = stats_raw.get("pair_bootstrap_cis", {})
    cl_analysis  = load_json("cl_analysis.json")
    ctx_dist_all = cl_analysis.get("context_distribution", {})

    try:
        pair_analysis = load_json("pair_analysis.json")
    except Exception:
        pair_analysis = {}

    try:
        chart_annotations = load_json("chart_annotations.json")
    except Exception:
        chart_annotations = {}

    all_pairs = manifest["pairs"]

    # Sort: pair 1 first, then by total descending
    pair1 = [p for p in all_pairs if p["id"] == 1]
    others = sorted([p for p in all_pairs if p["id"] != 1],
                    key=lambda p: p.get("total", 0) or 0, reverse=True)
    sorted_pairs = pair1 + others

    print(f"Generating DOCX for {len(sorted_pairs)} pairs...")

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

    # ── Default font ───────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    # ── Title page ─────────────────────────────────────────────────────
    doc.add_paragraph("")  # spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Toponym Pair Audit Reports")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE_UA

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("59 Ukrainian Toponym Spelling Transitions\nData · AI · Computational Linguistics")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_mentions = sum(p.get("total", 0) or 0 for p in all_pairs)
    run = meta.add_run(
        f"{manifest['total_pairs']} pairs · {manifest['records_scanned']} records scanned · "
        f"{fmt_number(total_mentions)} toponym mentions\n"
        f"8 data sources · 72 LLM models · {manifest['time_span']}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("kyivnotkiev.org")
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_UA
    run.bold = True

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────────
    toc_title = doc.add_paragraph()
    run = toc_title.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK

    doc.add_paragraph("")

    for i, pair in enumerate(sorted_pairs):
        pid = pair["id"]
        ru, ua = pair["russian"], pair["ukrainian"]
        adoption = pair.get("adoption", 0) or 0
        total = pair.get("total", 0) or 0
        cat = pair["category"]
        toc_entry = doc.add_paragraph()
        run = toc_entry.add_run(f"{i+1}. ")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run = toc_entry.add_run(f"{ru} → {ua}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run = toc_entry.add_run(f"  —  {cat.title()}, {adoption:.1f}%, {fmt_number(total)} mentions")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    doc.add_page_break()

    # ── Generate each pair report ──────────────────────────────────────
    for pair_idx, pair in enumerate(sorted_pairs):
        pid = pair["id"]
        ru = pair["russian"]
        ua = pair["ukrainian"]
        cat = pair["category"]
        adoption = pair.get("adoption", 0) or 0
        total = pair.get("total", 0) or 0
        starred = pair.get("starred", False)
        starred_label = pair.get("starred_label", "")

        ts = timeseries.get(str(pid), {})
        llm = llm_per_pair.get(str(pid), {})
        llm_summary = llm.get("summary", {})
        coll = collocations.get(str(pid), {})
        stat = stats.get(str(pid), {})
        ctx_dist = ctx_dist_all.get(str(pid), {})
        p_events = pair_events.get(str(pid), [])

        tas = llm_summary.get("tas_mean", 0) or 0
        recall = llm_summary.get("open_pct", 0) or 0
        forced_ru = llm_summary.get("forced_ru_pct", 0) or 0
        forced_ua = llm_summary.get("forced_ua_pct", 0) or 0
        forced = (forced_ru + forced_ua) / 2

        print(f"  [{pair_idx+1}/{len(sorted_pairs)}] Pair {pid}: {ru} → {ua}")

        # ── Section title ──────────────────────────────────────────────
        heading = doc.add_paragraph()
        run = heading.add_run(f"{ru}")
        run.font.size = Pt(20)
        run.font.color.rgb = RED_RU
        run.bold = True
        run = heading.add_run(" → ")
        run.font.size = Pt(20)
        run.font.color.rgb = GRAY
        run = heading.add_run(f"{ua}")
        run.font.size = Pt(20)
        run.font.color.rgb = BLUE_UA
        run.bold = True

        # Subtitle
        star_str = f" · {starred_label}" if starred else ""
        sub = doc.add_paragraph()
        run = sub.add_run(f"{cat.title()}{star_str}")
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run.italic = True

        # ── Stats line ─────────────────────────────────────────────────
        stats_p = doc.add_paragraph()
        run = stats_p.add_run(f"Adoption: {adoption:.1f}%")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE_UA
        run = stats_p.add_run(f"    |    Total Mentions: {fmt_number(total)}")
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        if tas > 0:
            run = stats_p.add_run(f"    |    TAS: {tas:.1f}%")
            run.font.size = Pt(11)
            run.font.color.rgb = PURPLE

        ci_lo = stat.get("lo", 0)
        ci_hi = stat.get("hi", 0)
        if ci_lo and ci_hi:
            ci_p = doc.add_paragraph()
            run = ci_p.add_run(f"Bootstrap 95% CI: [{ci_lo:.2f}%, {ci_hi:.2f}%], n = {fmt_number(stat.get('n', 0))}")
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY

        # ── Analysis text ──────────────────────────────────────────────
        analysis_heading = doc.add_paragraph()
        run = analysis_heading.add_run("Analysis")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK

        analysis_html = pair_analysis.get(str(pid), "")
        if analysis_html:
            analysis_text = strip_html(analysis_html)
            # Split by double newlines to create separate paragraphs
            paragraphs = [p.strip() for p in analysis_text.split('\n\n') if p.strip()]
            if not paragraphs:
                paragraphs = [analysis_text]
            for para_text in paragraphs:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK
        else:
            p = doc.add_paragraph()
            run = p.add_run("No hand-crafted analysis available for this pair.")
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.italic = True

        # ── Source data table ──────────────────────────────────────────
        src_heading = doc.add_paragraph()
        run = src_heading.add_run("Data Sources")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK

        # Build source data
        source_rows = []
        chart_jobs = {}  # key -> svg string for batch conversion

        for src in SOURCE_ORDER:
            sd = ts.get(src, [])
            if not sd:
                continue
            name = SOURCE_NAMES.get(src, src)
            total_src = sum((d.get("rus", 0) or 0) + (d.get("ukr", 0) or 0) for d in sd)
            adoptions = [d.get("adoption", 0) or 0 for d in sd]
            alltime = sum(adoptions) / len(adoptions) if adoptions else 0
            last12 = sd[-12:] if len(sd) >= 12 else sd
            last12_avg = sum(d.get("adoption", 0) or 0 for d in last12) / len(last12) if last12 else 0
            source_rows.append((name, total_src, alltime, last12_avg, src))

            # Build SVG chart for this source
            color = SOURCE_COLORS.get(src, "#6b7280")
            dates = [d.get("date", "") for d in sd if d.get("date")]
            start_yr = int(dates[0][:4]) if dates else 2010
            end_yr = int(dates[-1][:4]) if dates else 2026
            pair_annos = chart_annotations.get(str(pid), {}).get(src, [])
            svg = make_svg_chart(sd, color, start_yr, end_yr, p_events, pair_annos)
            if svg and '<svg' in svg:
                chart_jobs[f"{pid}_{src}"] = svg

        if source_rows:
            tbl = doc.add_table(rows=1, cols=4)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tbl)
            make_header_row(tbl, ["Source", "Total Mentions", "All-time %", "Last 12mo %"])
            R = WD_ALIGN_PARAGRAPH.RIGHT
            L = WD_ALIGN_PARAGRAPH.LEFT
            for name, total_src, alltime, last12_avg, _ in source_rows:
                add_data_row(tbl, [name, fmt_number(total_src), f"{alltime:.1f}%", f"{last12_avg:.1f}%"],
                             alignments=[L, R, R, R], bold_cols={3})

        # ── Convert and embed SVG charts ───────────────────────────────
        if chart_jobs:
            png_results = convert_chart_batch(chart_jobs)
            for src in SOURCE_ORDER:
                key = f"{pid}_{src}"
                png_bytes = png_results.get(key)
                if png_bytes:
                    name = SOURCE_NAMES.get(src, src)
                    chart_label = doc.add_paragraph()
                    run = chart_label.add_run(f"{name}")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK
                    stream = io.BytesIO(png_bytes)
                    doc.add_picture(stream, width=Inches(6.2))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── LLM Audit ─────────────────────────────────────────────────
        llm_models = llm.get("models", [])
        if tas > 0 and llm_models:
            llm_heading = doc.add_paragraph()
            run = llm_heading.add_run("AI / LLM Audit")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DARK

            gap = forced - recall
            n_models = llm_summary.get("n_models", 72)
            summary_p = doc.add_paragraph()
            run = summary_p.add_run(
                f"{n_models} models, 8 families. "
                f"Free Recall: {recall:.1f}%, Forced Choice: {forced:.1f}%, TAS: {tas:.1f}%. "
                f"Recognition-recall gap: {abs(gap):.0f}pp."
            )
            run.font.size = Pt(9)
            run.font.color.rgb = DARK

            # Best per family table
            best_per_fam = {}
            for m in llm_models:
                fam = m["family"]
                mtas = m.get("tas", 0) or 0
                if fam not in best_per_fam or mtas > (best_per_fam[fam].get("tas", 0) or 0):
                    best_per_fam[fam] = m

            if best_per_fam:
                tbl = doc.add_table(rows=1, cols=5)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl)
                make_header_row(tbl, ["Family", "Best Model", "Recall", "Forced", "TAS"])
                R = WD_ALIGN_PARAGRAPH.RIGHT
                C = WD_ALIGN_PARAGRAPH.CENTER
                L = WD_ALIGN_PARAGRAPH.LEFT
                for fam in FAM_ORDER:
                    if fam not in best_per_fam:
                        continue
                    m = best_per_fam[fam]
                    mtas = m.get("tas", 0) or 0
                    rec = "Yes" if m.get("x_open", 0) else "No"
                    fru = m.get("x_forced_ru_first", 0) or 0
                    fua = m.get("x_forced_ua_first", 0) or 0
                    favg = (fru + fua) / 2
                    forced_str = "Yes" if favg >= 0.75 else ("~" if favg > 0 else "No")
                    fam_short = fam.split()[-1] if " " in fam else fam
                    add_data_row(tbl, [fam_short, m["key"], rec, forced_str, f"{mtas*100:.0f}%"],
                                 alignments=[L, L, C, C, R], bold_cols={4})

        # ── Context Distribution ───────────────────────────────────────
        ru_ctx = ctx_dist.get("russian", {})
        ua_ctx = ctx_dist.get("ukrainian", {})
        if ru_ctx or ua_ctx:
            ctx_heading = doc.add_paragraph()
            run = ctx_heading.add_run("Context Distribution")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DARK

            coll_total = coll.get("total", 0)
            ctx_p = doc.add_paragraph()
            run = ctx_p.add_run(f"DeBERTa-v3-large classifier, F1=85.7%, {coll_total} texts")
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY

            all_contexts = set(list(ru_ctx.keys()) + list(ua_ctx.keys()))
            ctx_items = [(cn, ru_ctx.get(cn, 0) or 0, ua_ctx.get(cn, 0) or 0) for cn in all_contexts]
            ctx_items.sort(key=lambda x: x[1] + x[2], reverse=True)

            # Smart bucketing: 4%+, 2%+, 1%+, then all — max 10
            selected = []
            for threshold in [0.04, 0.02, 0.01, 0]:
                selected = [(cn, rp, up) for cn, rp, up in ctx_items if rp + up > threshold]
                if len(selected) >= 4:
                    break
            selected = selected[:10]

            if selected:
                tbl = doc.add_table(rows=1, cols=3)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl)
                make_header_row(tbl, ["Context", f"RU % ({ru})", f"UA % ({ua})"])
                R = WD_ALIGN_PARAGRAPH.RIGHT
                L = WD_ALIGN_PARAGRAPH.LEFT
                for cn, rp, up in selected:
                    total_p = rp + up
                    if total_p < 0.005:
                        continue
                    ru_pct = rp / total_p * 100
                    ua_pct = up / total_p * 100
                    label = CTX_LABELS.get(cn, cn.replace("_", " ").title())
                    add_data_row(tbl, [label, f"{ru_pct:.0f}%", f"{ua_pct:.0f}%"],
                                 alignments=[L, R, R])

        # ── Collocations ───────────────────────────────────────────────
        ru_coll = coll.get("russian", {}).get("collocates", []) if isinstance(coll.get("russian"), dict) else []
        ua_coll = coll.get("ukrainian", {}).get("collocates", []) if isinstance(coll.get("ukrainian"), dict) else []

        if ru_coll or ua_coll:
            coll_heading = doc.add_paragraph()
            run = coll_heading.add_run("Contrastive Analysis (log-odds ratio)")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DARK

            if ru_coll:
                p = doc.add_paragraph()
                run = p.add_run(f"{ru} form: ")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RED_RU
                words = ", ".join(c["word"] for c in ru_coll[:10])
                run = p.add_run(words)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

            if ua_coll:
                p = doc.add_paragraph()
                run = p.add_run(f"{ua} form: ")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = BLUE_UA
                words = ", ".join(c["word"] for c in ua_coll[:10])
                run = p.add_run(words)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

        # ── Religious Sources ──────────────────────────────────────────
        rel_rows = []
        for denom in religious.get("denominations", []):
            for pp in denom.get("pairs", []):
                if pp.get("pair_id") == pid:
                    rc = pp.get("ru", 0)
                    uc = pp.get("ua", 0)
                    if rc + uc == 0:
                        continue
                    ua_pct = pp.get("ua_pct", 0) or 0
                    dn = denom["name"]
                    short = DENOM_SHORT.get(dn, dn)
                    rel_rows.append((short, rc, uc, ua_pct))

        if rel_rows:
            rel_heading = doc.add_paragraph()
            run = rel_heading.add_run("Religious Sources")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DARK

            tbl = doc.add_table(rows=1, cols=4)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tbl)
            make_header_row(tbl, ["Institution", "Russian form", "Ukrainian form", "UA %"])
            R = WD_ALIGN_PARAGRAPH.RIGHT
            L = WD_ALIGN_PARAGRAPH.LEFT
            for short, rc, uc, ua_pct in rel_rows:
                add_data_row(tbl, [short, str(rc), str(uc), f"{ua_pct:.1f}%"],
                             alignments=[L, R, R, R], bold_cols={3})

        # ── References ─────────────────────────────────────────────────
        ref_heading = doc.add_paragraph()
        run = ref_heading.add_run("References")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK

        events_list = [
            "2018-10: #KyivNotKiev campaign launched by MFA of Ukraine",
            "2019-06: US Board on Geographic Names adopts Kyiv",
            "2019-08: AP Stylebook adopts Kyiv",
            "2022-02: Full-scale Russian invasion accelerates global adoption",
        ]
        for ev in p_events:
            events_list.append(f"{ev['date']}: {ev['label']}")

        ref_p = doc.add_paragraph()
        run = ref_p.add_run("Key Events: ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        run = ref_p.add_run(" | ".join(events_list))
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

        ref_p2 = doc.add_paragraph()
        run = ref_p2.add_run("Data Sources: ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        run = ref_p2.add_run(
            "Google Trends (55 countries) · GDELT (53K+ domains) · Wikipedia Pageviews · "
            "Reddit (Arctic Shift) · YouTube (yt-dlp) · Google Books Ngrams (corpus 37) · "
            "OpenAlex (250M+ works) · Open Library (8M+ books)"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

        # ── Page break between pairs ───────────────────────────────────
        if pair_idx < len(sorted_pairs) - 1:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────
    print(f"\nSaving to {OUT_FILE}...")
    doc.save(str(OUT_FILE))
    size_mb = OUT_FILE.stat().st_size / (1024 * 1024)
    print(f"Done! {OUT_FILE} ({size_mb:.1f} MB)")


if __name__ == "__main__":
    build_docx()
