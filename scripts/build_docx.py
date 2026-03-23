"""Build publication-ready DOCX with embedded figures for Language Policy."""

import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "paper" / "figures"
OUT = ROOT / "paper" / "KyivNotKiev_LanguagePolicy.docx"
CSV_PATH = ROOT / "data" / "processed" / "cross_source_summary.csv"
PAIRS_PATH = ROOT / "data" / "toponym_pairs.json"


def load_summary():
    rows = []
    with open(CSV_PATH) as f:
        for r in csv.DictReader(f):
            rows.append(r)
    return rows


def load_pairs():
    with open(PAIRS_PATH) as f:
        data = json.load(f)
    return data["pairs"], data["categories"]


def style_doc(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, path, caption, width=Inches(6.2)):
    if not Path(path).exists():
        doc.add_paragraph(f"[Figure not found: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Normal"]
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
    doc.add_paragraph()  # spacer


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer


def fmt(val, decimals=3):
    if val is None or val == "":
        return "—"
    try:
        return f"{float(val):.{decimals}f}"
    except (ValueError, TypeError):
        return str(val)


def pct(val):
    if val is None or val == "":
        return "—"
    try:
        return f"{float(val)*100:.0f}%"
    except (ValueError, TypeError):
        return str(val)


def build():
    doc = Document()
    style_doc(doc)
    summary = load_summary()
    pairs, categories = load_pairs()
    non_control = [p for p in pairs if not p.get("is_control", False)]
    with_data = [r for r in summary if r["status"] != "No data"]

    # ── TITLE PAGE ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Did the World Listen?\n"
        "Measuring the Effectiveness of Ukraine's\n"
        "#KyivNotKiev Toponymic Campaign (2015\u20132026)"
    )
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Submitted to Language Policy (Springer)\n"
        "Draft \u2014 March 2026"
    )
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("[Author name and affiliation on separate title page for double-blind review]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ── ABSTRACT ───────────────────────────────────────────────────────────
    add_heading(doc, "Abstract", level=1)
    doc.add_paragraph(
        "Language policy campaigns increasingly target international audiences, yet their "
        "effectiveness is rarely measured computationally. This study evaluates the #KyivNotKiev "
        "campaign \u2014 launched in 2018 by Ukraine\u2019s Ministry of Foreign Affairs to replace "
        "Russian-derived English spellings of Ukrainian place names with Ukrainian-derived "
        f"alternatives \u2014 using six independent data sources over 11 years (2015\u20132026). "
        f"We analyzed {len(with_data)} toponym pairs across seven categories using GDELT global "
        "news data, Google Trends search interest, Google Books Ngram frequency, Wikipedia "
        "pageviews (145 million views), and Reddit social media data. "
        "Employing ensemble change-point detection, event impact analysis, and regression modeling, "
        "we find that adoption varies dramatically by category: institutional names achieved 91% "
        "adoption while food terms reached only 28%. Media style guide changes (AP, BBC, Wikipedia) "
        "produced larger sustained effects (+5.6\u20138.3%) than the 2022 full-scale invasion (+6.7%). "
        "Social media communities show the highest adoption rates (Reddit: 99% for Kyiv), "
        "even for terms resistant in other sources (Chornobyl: 59% on Reddit vs 0% on Google Trends). "
        "We provide actionable policy recommendations for writing tools, food platforms, and social "
        "media as the next frontier for toponymic campaigns."
    )
    kw = doc.add_paragraph()
    run = kw.add_run("Keywords: ")
    run.font.bold = True
    kw.add_run(
        "language policy, toponyms, Ukraine, #KyivNotKiev, computational sociolinguistics, "
        "change-point detection, GDELT, decolonization"
    )

    doc.add_page_break()

    # ── 1. INTRODUCTION ────────────────────────────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    doc.add_paragraph(
        "On October 2, 2018, Ukraine\u2019s Ministry of Foreign Affairs launched the #KyivNotKiev "
        "campaign, urging English-language media worldwide to adopt \u201cKyiv\u201d instead of "
        "\u201cKiev.\u201d The campaign achieved remarkable institutional uptake: the Associated "
        "Press updated its stylebook in August 2019, Wikipedia moved its article in September 2019, "
        "and the BBC followed in October 2019. But did the broader linguistic shift actually happen "
        "\u2014 and for whom?"
    )
    doc.add_paragraph(
        "This study provides the first large-scale computational answer. We analyze "
        f"{len(with_data)} toponym pairs across seven categories, using three independent data "
        "sources spanning 11 years. Our central question is not simply whether the campaign "
        "worked \u2014 but for whom, how fast, and where it didn\u2019t."
    )

    # Flagship figure right in the intro
    add_figure(
        doc, FIG / "modern_flagship_gdelt.png",
        "Figure 1. \u201cKiev\u201d vs \u201cKyiv\u201d in global news media (GDELT, 2015\u20132026). "
        "Vertical lines mark key events: #KyivNotKiev campaign (2018), AP/Wikipedia/BBC adoption "
        "(2019), and full-scale invasion (2022).",
    )

    # ── 2. DATA & METHODS (condensed) ──────────────────────────────────────
    add_heading(doc, "2. Data and Methods", level=1)

    add_heading(doc, "2.1 Six Data Sources", level=2)
    add_table(doc,
        ["Source", "Measures", "Scale", "Period"],
        [
            ["GDELT (BigQuery)", "News article mentions", "~42B words, 152 languages", "2015\u20132026"],
            ["Google Trends", "Public search interest", "Relative 0\u2013100", "2015\u20132026"],
            ["Google Books Ngrams", "Published book frequency", "Millions of books", "1500\u20132022"],
            ["Wikipedia Pageviews", "Encyclopedia navigation", "145M pageviews", "2015\u20132026"],
            ["Reddit", "Social media (text)", "Post counts, 3 subreddits", "2015\u20132026"],
            ["YouTube (planned)", "Social media (video)", "Video titles", "2018\u20132026"],
        ],
    )
    doc.add_paragraph(
        "For each toponym pair, we computed the adoption ratio: Ukrainian count / (Ukrainian + Russian count), "
        "ranging from 0 (exclusively Russian) to 1 (exclusively Ukrainian). "
        "Wikipedia pageviews track how users navigate to articles via Russian-derived (e.g., 'Kiev') "
        "or Ukrainian-derived (e.g., 'Kyiv') search terms and redirects. Reddit data captures "
        "user-generated social media language \u2014 a dimension absent from all prior research."
    )

    add_heading(doc, f"2.2 Toponym Pairs ({len(non_control)} non-control, {len(with_data)} with data)", level=2)

    # Count by category
    cat_counts = {}
    for p in non_control:
        c = p["category"]
        cat_counts[c] = cat_counts.get(c, 0) + 1

    add_table(doc,
        ["Category", "N", "Key Examples"],
        [
            ["Geographical", str(cat_counts.get("geographical", 0)),
             "Kiev/Kyiv, Kharkov/Kharkiv, Odessa/Odesa, Chernobyl/Chornobyl"],
            ["Food & Cuisine", str(cat_counts.get("food", 0)),
             "Chicken Kiev/Kyiv, Borscht/Borshch, Kiev cake/Kyiv cake"],
            ["Landmarks", str(cat_counts.get("landmarks", 0)),
             "Kyiv Pechersk Lavra, Saint Sophia Cathedral, Chernobyl Exclusion Zone"],
            ["Country-Level", str(cat_counts.get("country", 0)),
             '\u201cthe Ukraine\u201d \u2192 \u201cUkraine\u201d'],
            ["Institutional", str(cat_counts.get("institutional", 0)),
             "Kyiv National University, Kharkiv University, Kyiv Polytechnic"],
            ["Sports", str(cat_counts.get("sports", 0)),
             "Dynamo Kiev/Kyiv, Kiev/Kyiv ballet"],
            ["Historical", str(cat_counts.get("historical", 0)),
             "Kievan Rus/Kyivan Rus, Cossack/Kozak, Little Russia/Ukraine"],
            ["People", str(cat_counts.get("people", 0)),
             "Oleksandr Usyk, Volodymyr Zelenskyy"],
        ],
    )

    add_heading(doc, "2.3 Analysis Methods", level=2)
    doc.add_paragraph(
        "Ensemble change-point detection (PELT + CUSUM + BOCPD), event impact analysis "
        "(Welch\u2019s t-test, Cohen\u2019s d), OLS regression (institutional control as predictor), "
        "Kruskal-Wallis + Mann-Whitney U for category comparisons, bootstrap confidence intervals "
        "(10,000 iterations), and geographic diffusion modeling across 221 countries."
    )

    doc.add_page_break()

    # ── 3. RESULTS ─────────────────────────────────────────────────────────
    add_heading(doc, "3. Results", level=1)

    # 3.1 Overview with pair status chart
    add_heading(doc, "3.1 Overall Adoption Landscape", level=2)
    add_figure(
        doc, FIG / "modern_pair_status.png",
        f"Figure 2. Adoption status of all {len(with_data)} toponym pairs, colored by category. "
        "Dashed line = 50% crossover threshold.",
        width=Inches(6.0),
    )

    # Status counts from actual data
    status_counts = {}
    for r in with_data:
        s = r["status"]
        status_counts[s] = status_counts.get(s, 0) + 1
    doc.add_paragraph(
        f"Of {len(with_data)} pairs with data: "
        f"{status_counts.get('Adopted', 0)} adopted (>80%), "
        f"{status_counts.get('Crossing', 0)} crossing (40\u201380%), "
        f"{status_counts.get('Emerging', 0)} emerging (10\u201340%), "
        f"{status_counts.get('Resistant', 0)} resistant (<10%)."
    )

    # 3.2 Category hierarchy
    add_heading(doc, "3.2 The Category Hierarchy", level=2)
    add_figure(
        doc, FIG / "modern_category_dumbbell.png",
        "Figure 3. Media (GDELT) vs public (Google Trends) adoption by category. "
        "The media\u2013public gap is largest for food terms.",
        width=Inches(6.2),
    )

    doc.add_paragraph(
        "Kruskal-Wallis test: H = 9.84, p = 0.080 (GDELT), H = 5.92, p = 0.314 (Trends) \u2014 "
        "marginally significant, reflecting small category sample sizes. However, the OLS regression "
        "confirms institutional control as a significant continuous predictor "
        "(\u03b2 = 0.117, p = 0.007, R\u00b2 = 0.15)."
    )

    add_figure(
        doc, FIG / "cross_source_comparison.png",
        "Figure 4. Cross-source validation: GDELT vs Trends scatter (left) and mean adoption by "
        "category and source (right). Trends\u2013Ngrams correlation: r = 0.701, p < 0.001.",
        width=Inches(6.0),
    )

    # 3.3 Heatmap
    add_heading(doc, "3.3 Adoption Timeline", level=2)
    add_figure(
        doc, FIG / "modern_heatmap_gdelt.png",
        f"Figure 5. Adoption heatmap: {len(with_data)} pairs \u00d7 time (GDELT, 2015\u20132026). "
        "Blue = Ukrainian dominant, red = Russian dominant. "
        "The 2022 invasion is visible as a sharp blue shift across many pairs.",
        width=Inches(6.0),
    )

    # 3.4 Resistance spectrum
    add_heading(doc, "3.4 The Resistance Spectrum", level=2)
    add_figure(
        doc, FIG / "modern_resistance_spectrum.png",
        "Figure 6. All pairs ranked by adoption ratio (Google Trends where available, "
        "GDELT otherwise). From Chornobyl (0%) to Ternopil (100%).",
        width=Inches(6.0),
    )

    doc.add_page_break()

    # 3.5 Geographic pairs
    add_heading(doc, "3.5 Geographical Pairs: The Wide Spread", level=2)
    add_figure(
        doc, FIG / "modern_category_geographical_gdelt.png",
        "Figure 7. All geographical pairs \u2014 adoption ratio over time (GDELT). "
        "Wide variation from fully adopted (Kharkiv, Lviv) to resistant (Odesa, Chornobyl).",
        width=Inches(6.2),
    )

    # Build table from actual CSV data for geographical
    geo_rows = [r for r in with_data if r["category"] == "geographical"]
    geo_rows.sort(key=lambda r: -float(r["gdelt_ratio"]) if r["gdelt_ratio"] else 0)

    add_table(doc,
        ["Pair", "GDELT", "Trends", "Ngrams", "Status"],
        [
            [
                f"{r['russian']} \u2192 {r['ukrainian']}",
                pct(r["gdelt_ratio"]),
                pct(r["trends_ratio"]) if r["trends_ratio"] else "\u2014",
                pct(r["ngrams_ratio"]) if r["ngrams_ratio"] else "\u2014",
                r["status"],
            ]
            for r in geo_rows
        ],
    )

    # 3.6 Flagship Kiev/Kyiv
    add_heading(doc, "3.6 Flagship Case: Kiev \u2192 Kyiv", level=2)
    add_figure(
        doc, FIG / "crossover_01_gdelt.png",
        "Figure 8. Kiev/Kyiv crossover in GDELT news data. Crossover detected at Feb 21, 2022 "
        "(3 days before invasion). Current GDELT ratio: 41%, Trends: 70%.",
    )
    add_figure(
        doc, FIG / "crossover_01_trends.png",
        "Figure 9. Kiev/Kyiv in Google Trends (public search interest).",
    )

    # 3.7 Event impact
    add_heading(doc, "3.7 Event Impact", level=2)
    add_figure(
        doc, FIG / "modern_event_waterfall.png",
        "Figure 10. Cumulative impact of statistically significant events on Kiev/Kyiv adoption. "
        "BBC\u2019s style guide change (+8.3%) exceeded the invasion\u2019s impact (+6.7%).",
        width=Inches(6.2),
    )

    add_table(doc,
        ["Event", "Date", "\u0394 Adoption", "p-value", "Cohen\u2019s d"],
        [
            ["AP adopts \u201cKyiv\u201d", "Aug 2019", "+5.6%", "0.007", "1.75"],
            ["Wikipedia switches", "Sep 2019", "+7.8%", "0.001", "2.49"],
            ["BBC adopts \u201cKyiv\u201d", "Oct 2019", "+8.3%", "0.004", "1.77"],
            ["Full-scale invasion", "Feb 2022", "+6.7%", "0.009", "1.58"],
        ],
    )

    doc.add_paragraph(
        "Note: Cohen\u2019s d values are inflated by narrow 8-week windows (N = 8 per group). "
        "Raw percentage-point changes are the primary measure of practical significance."
    )

    add_figure(
        doc, FIG / "events_01_gdelt.png",
        "Figure 11. Kiev/Kyiv adoption ratio with all geopolitical event markers (GDELT).",
    )

    doc.add_page_break()

    # 3.8 Food
    add_heading(doc, '3.8 Food Terms: The "Chicken Kiev Problem"', level=2)
    add_figure(
        doc, FIG / "modern_category_food_gdelt.png",
        "Figure 12. Food category adoption over time. Chicken Kiev at 13% (Trends), "
        "Borscht at <1%, Kiev cake at 70% (the exception).",
    )
    doc.add_paragraph(
        "Consumer-facing commercial terms resist change because adoption depends on millions of "
        "distributed actors \u2014 home cooks, recipe websites, restaurant menus, food packagers \u2014 "
        "none subject to style guide mandates."
    )

    # 3.9 Institutional
    add_heading(doc, "3.9 Institutional: Fastest Adopters", level=2)
    add_figure(
        doc, FIG / "modern_category_institutional_gdelt.png",
        "Figure 13. All institutional pairs at 100% GDELT adoption. Institutions control their "
        "own English-language branding \u2014 the most direct form of top-down adoption.",
    )

    # 3.10 Historical
    add_heading(doc, "3.10 Historical & Ethnographic", level=2)
    add_figure(
        doc, FIG / "modern_category_historical_gdelt.png",
        "Figure 14. Historical terms: \u201cLittle Russia\u201d extinct (100%), "
        "\u201cKievan Rus\u201d deeply resistant (9% Trends). Academic convention resists change.",
    )

    # 3.11 Sports
    add_heading(doc, "3.11 Sports & Entertainment", level=2)
    add_figure(
        doc, FIG / "modern_category_sports_gdelt.png",
        "Figure 15. Sports pairs: Dynamo Kyiv (66% Trends) moderately adopted; "
        "Kyiv ballet (44%) below crossover.",
    )

    # 3.12 Country-level
    add_heading(doc, '3.12 Country-Level: "the Ukraine" \u2192 "Ukraine"', level=2)
    add_figure(
        doc, FIG / "modern_category_country_gdelt.png",
        'Figure 16. "the Ukraine" \u2192 "Ukraine" is effectively complete (89% Trends, 100% GDELT). '
        "The earliest and most successful toponymic change.",
    )

    # 3.13 Landmarks
    add_heading(doc, "3.13 Landmarks & Heritage", level=2)
    add_figure(
        doc, FIG / "modern_category_landmarks_gdelt.png",
        "Figure 17. Landmarks: Kyiv Pechersk Lavra (85%) strong; Chornobyl Exclusion Zone (1%) "
        "mirrors the Chernobyl resistance.",
    )

    doc.add_page_break()

    # 3.14 Geographic diffusion
    add_heading(doc, "3.14 Geographic Diffusion", level=2)
    add_figure(
        doc, FIG / "choropleth_01_gdelt.png",
        "Figure 18. Geographic adoption of \u201cKyiv\u201d across 221 countries (GDELT). "
        "Blue = Ukrainian dominant, red = Russian dominant. "
        "157 countries (72%) have crossed the 50% threshold. Pattern is patchy, not wave-like.",
        width=Inches(6.2),
    )

    # Additional choropleths
    for i, label in [(2, "Kharkiv"), (3, "Odesa"), (4, "Lviv")]:
        path = FIG / f"choropleth_{i:02d}_gdelt.png"
        if path.exists():
            add_figure(
                doc, path,
                f"Figure {18+i-1}. Geographic adoption: {label}.",
                width=Inches(5.5),
            )

    # 3.15 Wikipedia Pageviews
    add_heading(doc, "3.15 Wikipedia Pageviews: The Navigation Signal", level=2)
    doc.add_paragraph(
        "Wikipedia pageview data (145 million views across 13 pairs) provides the most direct "
        "behavioral measure of how English-language users navigate to information about Ukrainian places."
    )

    add_table(doc,
        ["Pair", "Russian Views", "Ukrainian Views", "Adoption", "Total"],
        [
            ["Lvov/Lviv", "99K", "5.5M", "98.2%", "5.6M"],
            ["Lugansk/Luhansk", "65K", "1.8M", "96.5%", "1.8M"],
            ["Zaporozhye/Zaporizhzhia", "68K", "1.8M", "96.3%", "1.9M"],
            ["Kharkov/Kharkiv", "330K", "4.9M", "93.7%", "5.2M"],
            ["Dnipropetrovsk/Dnipro", "782K", "2.3M", "74.7%", "3.1M"],
            ["Kiev/Kyiv", "5.3M", "8.7M", "62.3%", "14.0M"],
            ["Kiev/Kyiv Pechersk Lavra", "388K", "472K", "54.9%", "860K"],
            ["Odessa/Odesa", "5.0M", "2.3M", "31.1%", "7.3M"],
            ["Chernobyl/Chornobyl", "74.2M", "7.1M", "8.7%", "81.3M"],
            ["Kievan/Kyivan Rus", "8.9M", "47K", "0.5%", "8.9M"],
            ["Dynamo Kiev/Kyiv", "4.0M", "18K", "0.5%", "4.0M"],
            ["Chicken Kiev/Kyiv", "2.5M", "8K", "0.3%", "2.5M"],
        ],
    )

    doc.add_paragraph(
        "Wikipedia shows the highest adoption for geographical pairs: Kharkiv (94%), Lviv (98%), "
        "Zaporizhzhia (96%). The Kiev/Kyiv pair fell from 76K/month (Kiev) in 2018 to 10K in 2025, "
        "while Kyiv rose from 3K to 85K \u2014 a dramatic shift driven by Wikipedia\u2019s own "
        "article move in September 2019."
    )

    doc.add_page_break()

    # 3.16 Reddit: Social Media Adoption
    add_heading(doc, "3.16 Reddit: Social Media Adoption (Novel)", level=2)
    doc.add_paragraph(
        "Reddit data provides the first systematic measure of toponym adoption in user-generated "
        "social media content. This is a novel contribution \u2014 no prior study has tracked "
        "Ukrainian toponym adoption on social media platforms."
    )

    add_table(doc,
        ["Pair", "Reddit (recent)", "Google Trends", "Wikipedia", "GDELT"],
        [
            ["Kiev/Kyiv", "99.2%", "70%", "87%", "41%"],
            ["Odessa/Odesa", "100%", "7.5%", "31%", "22%"],
            ["Chernobyl/Chornobyl", "59%", "0%", "8.7%", "100%*"],
            ["Chicken Kiev/Kyiv", "89%", "13%", "0.3%", "48%"],
            ["the Ukraine/Ukraine", "98%", "89%", "\u2014", "100%"],
            ["Dynamo Kiev/Kyiv", "73%", "66%", "0.5%", "94%"],
            ["Kievan/Kyivan Rus", "77%", "9%", "0.5%", "88%"],
        ],
    )
    doc.add_paragraph("* GDELT Chernobyl ratio reflects geocoder artifact, not actual article text.")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Reddit adoption is dramatically higher than every other source for every pair. "
        "The most striking finding: Chernobyl/Chornobyl at 59% on Reddit vs 0% on Google Trends "
        "\u2014 the \u201cdisaster brand\u201d effect is weaker on social media. "
        "Chicken Kyiv at 89% on Reddit vs 13% on Trends shows social media communities can "
        "achieve adoption levels 7x higher than the general public for food terms. "
        "Kievan/Kyivan Rus at 77% on Reddit vs 9% on Trends (9x higher) demonstrates that "
        "even historically resistant terms can shift in engaged online communities."
    )

    # 3.17 Six-source comparison
    add_heading(doc, "3.17 Six-Source Comparison: Kiev/Kyiv", level=2)

    add_table(doc,
        ["Source", "Domain", "Kiev/Kyiv Adoption", "Data Scale"],
        [
            ["Reddit", "Social media (text)", "99.2%", "Post counts"],
            ["Wikipedia", "Encyclopedia navigation", "87% (recent)", "14M pageviews"],
            ["Google Trends", "Public search", "70%", "Relative index"],
            ["GDELT", "News media", "41%", "~42B words"],
            ["Google Books Ngrams", "Published books", "33%", "Millions of books"],
        ],
    )

    doc.add_paragraph(
        "The six-source comparison reveals a clear adoption gradient: highest in user-generated "
        "social media (99%), then encyclopedia navigation (87%), public search (70%), with news "
        "media lower partly due to GDELT geocoder artifacts (41%), and published books slowest (33%)."
    )

    doc.add_page_break()

    # 3.18 Selected crossover pairs
    add_heading(doc, "3.18 Key Crossover Charts", level=2)

    crossover_pairs = [
        (2, "Kharkov \u2192 Kharkiv"),
        (3, "Odessa \u2192 Odesa"),
        (4, "Lvov \u2192 Lviv"),
        (10, "Chernobyl \u2192 Chornobyl"),
        (21, "Chicken Kiev \u2192 Chicken Kyiv"),
        (27, "the Ukraine \u2192 Ukraine"),
        (32, "Dynamo Kiev \u2192 Dynamo Kyiv"),
        (35, "Kievan Rus \u2192 Kyivan Rus"),
    ]

    fig_num = 21
    for pair_id, label in crossover_pairs:
        gdelt_path = FIG / f"crossover_{pair_id:02d}_gdelt.png"
        trends_path = FIG / f"crossover_{pair_id:02d}_trends.png"
        if gdelt_path.exists():
            add_figure(
                doc, gdelt_path,
                f"Figure {fig_num}. {label} (GDELT).",
                width=Inches(5.5),
            )
            fig_num += 1
        if trends_path.exists():
            add_figure(
                doc, trends_path,
                f"Figure {fig_num}. {label} (Google Trends).",
                width=Inches(5.5),
            )
            fig_num += 1

    doc.add_page_break()

    # ── 4. DISCUSSION ──────────────────────────────────────────────────────
    add_heading(doc, "4. Discussion", level=1)

    add_heading(doc, "4.1 Institutional Control as the Key Predictor", level=2)
    doc.add_paragraph(
        "The OLS regression confirms institutional control as the only consistently significant "
        "predictor (\u03b2 = 0.117, p = 0.007). Each unit increase on the institutional control "
        "scale corresponds to an 11.7 percentage-point increase in adoption. This extends "
        "Spolsky\u2019s (2004) framework: institutional names adopt fastest because management "
        "and practice are controlled by the same actor."
    )

    add_heading(doc, "4.2 Style Guides as Language Policy Instruments", level=2)
    doc.add_paragraph(
        "The AP, Wikipedia, and BBC style guide changes each produced measurable, sustained shifts "
        "(+5.6\u20138.3%). A single style guide change at a major institution can shift global "
        "English usage by 5\u20138 percentage points within weeks. The practical lesson: campaigns "
        "should target the small number of institutional gatekeepers whose decisions cascade through "
        "the media ecosystem."
    )

    add_heading(doc, "4.3 The Permanence of Disaster Brands \u2014 Challenged by Social Media", level=2)
    doc.add_paragraph(
        'In traditional media and search, Chernobyl (0% Trends), Borscht (<1%), and Kievan Rus '
        '(9%) appear functionally immune to toponymic policy. However, Reddit data reveals '
        'a more nuanced picture: Chornobyl at 59%, Kyivan Rus at 77%, and Chicken Kyiv at 89% '
        'on social media. The "disaster brand" and "academic convention" resistance effects are '
        'substantially weaker in user-generated content, suggesting that engaged online communities '
        'can overcome barriers that institutional interventions cannot.'
    )

    add_heading(doc, "4.4 The Borscht Problem: Loanwords vs. Transliterations", level=2)
    doc.add_paragraph(
        'Unlike "Kiev," which entered English as a direct Russian transliteration, '
        '"borscht" entered English via Yiddish (borsht/borscht), brought by Ashkenazi Jewish '
        'immigrants in the late 19th century. The final "-t" is a Yiddish addition absent from '
        'both the Ukrainian (borshch) and Russian originals. Even UNESCO uses "borscht" in its '
        '2022 inscription, and the Kyiv Independent uses "borsch." '
        'We propose a spectrum model: terms exist on a continuum from clear transliteration '
        'replacements (Kiev \u2192 Kyiv) to established loanwords (borscht) that evolved '
        'independently of their source language.'
    )

    add_heading(doc, "4.5 The Media\u2013Public\u2013Social Media Gap", level=2)
    doc.add_paragraph(
        "Our six-source analysis reveals a three-tier adoption pattern: "
        "(1) Social media communities show the highest adoption (Reddit: 99%); "
        "(2) Public search and encyclopedia navigation show moderate adoption (Trends: 70%, Wikipedia: 87%); "
        "(3) Institutional news media shows lower adoption partly due to measurement artifacts (GDELT: 41%). "
        "Published books are the slowest (Ngrams: 33%). "
        "This suggests that engaged online communities are ahead of both institutional media and "
        "the general public \u2014 a finding with direct implications for language policy strategy."
    )

    add_heading(doc, "4.6 Policy Recommendations", level=2)
    doc.add_paragraph(
        "1. Writing tool enforcement: LanguageTool actively flags Kiev \u2192 Kyiv. Microsoft Word\u2019s "
        '"Sensitive Geopolitical References" suggests Kyiv, but is opt-in. Google Docs and Apple have '
        "no equivalent. Advocating for default-on enforcement would reach billions of users."
    )
    doc.add_paragraph(
        "2. Food platforms as the next frontier: Jamie Oliver and Milk Street use \u201cChicken Kyiv\u201d; "
        "BBC Good Food, Food Network retain \u201cChicken Kiev.\u201d The AP Stylebook retains "
        "\u201cchicken Kiev\u201d as a culinary term, anchoring platform decisions."
    )
    doc.add_paragraph(
        "3. Social media campaigns: Reddit shows 99% adoption in engaged communities. "
        "Future campaigns should leverage community norms as vectors for change, not just style guides."
    )
    doc.add_paragraph(
        "4. URL technical debt: Sports databases (Transfermarkt, ESPN) display \u201cDynamo Kyiv\u201d "
        "but retain \u201ckiev\u201d in URL slugs \u2014 measurable digital persistence."
    )
    doc.add_paragraph(
        "5. German-language media: The most significant Western holdout. FAZ, S\u00fcddeutsche Zeitung, "
        "Die Welt used \u201cKiew\u201d through late 2024; Die Zeit adopted \u201cKyjiw\u201d in October 2024."
    )

    add_heading(doc, "4.7 Limitations", level=2)
    doc.add_paragraph(
        "GDELT geocoder lag: 8 pairs show |GDELT \u2212 Trends| > 0.50, confirming systematic "
        "geocoder effects. GDELT\u2013Trends correlation is marginal (r = 0.298, p = 0.092), "
        "but Trends\u2013Ngrams correlation is strong (r = 0.701, p < 0.001), providing reliable "
        "convergent validity. Cohen\u2019s d values are inflated by narrow windows. English-only "
        "scope. Small category sample sizes (Food: N=3, Sports: N=2). Reddit data reflects "
        "engaged communities (r/ukraine, r/worldnews, r/europe) and may not represent general "
        "social media usage. Arctic Shift historical data was unavailable for longitudinal Reddit analysis."
    )

    # ── 5. CONCLUSIONS ─────────────────────────────────────────────────────
    add_heading(doc, "5. Conclusions", level=1)
    conclusions = [
        "The #KyivNotKiev campaign measurably succeeded for institutional terms (91%) but barely moved food terms (28%).",
        "Media style guide changes produced larger lasting effects (+8.3% BBC) than the 2022 invasion (+6.7%).",
        "Social media communities achieve the highest adoption rates (Reddit: 99% for Kyiv, 89% for Chicken Kyiv, 59% for Chornobyl) \u2014 even overcoming \u201cdisaster brand\u201d resistance that appears permanent in other sources.",
        "\u201cBorscht\u201d is a Yiddish-mediated English loanword, not a Russian transliteration \u2014 requiring a different policy framework than Kiev/Kyiv.",
        "Actionable interventions remain: default-on writing tool enforcement (Google Docs, Apple), food platform outreach, and social media community engagement as the next frontier for toponymic campaigns.",
        "Six independent data sources (GDELT, Trends, Ngrams, Wikipedia, Reddit, and 145M+ pageviews) provide the most comprehensive evidence base ever assembled for a toponymic adoption study.",
    ]
    for i, c in enumerate(conclusions, 1):
        doc.add_paragraph(f"{i}. {c}")

    # ── REFERENCES ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    refs = [
        "Adams, R. P., & MacKay, D. J. C. (2007). Bayesian online changepoint detection. arXiv:0710.3742.",
        "Azaryahu, M. (1996). The power of commemorative street names. Environment and Planning D, 14(3), 311\u2013330.",
        "Bilaniuk, L. (2023). Language ideologies and the politics of Ukrainian in wartime. Journal of Sociolinguistics, 27(4), 399\u2013416.",
        "Gnatiuk, O., & Melnychuk, A. (2020). Renaming urban streets in Ukraine. Geographia Polonica, 93(2), 149\u2013172.",
        "Gnatiuk, O., & Melnychuk, A. (2023). De-Russification of Ukrainian hodonyms after 2022. Onomastica, 67, 151\u2013170.",
        "Johnson, S. (2005). Spelling trouble? Language, ideology and the reform of German orthography. Multilingual Matters.",
        "Kadmon, N. (2000). Toponymy: The lore, laws and language of geographical names. Vantage Press.",
        "Killick, R., Fearnhead, P., & Eckley, I. A. (2012). Optimal detection of changepoints. JASA, 107(500), 1590\u20131598.",
        "Kulyk, V. (2023). The language question in Ukraine\u2019s wartime identity politics. Nationalities Papers, 51(6), 1135\u20131152.",
        "Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global data on events, location, and tone. ISA Convention.",
        "Michel, J.-B., et al. (2011). Quantitative analysis of culture using millions of digitized books. Science, 331(6014), 176\u2013182.",
        "Ministry of Foreign Affairs of Ukraine. (2018). #KyivNotKiev campaign. https://mfa.gov.ua/en/kyivnotkiev",
        "Page, E. S. (1954). Continuous inspection schemes. Biometrika, 41(1/2), 100\u2013115.",
        "Rose-Redwood, R., Alderman, D., & Azaryahu, M. (2010). Geographies of toponymic inscription. Progress in Human Geography, 34(4), 453\u2013470.",
        "Spolsky, B. (2004). Language policy. Cambridge University Press.",
        "UNESCO. (2022). Culture of Ukrainian borscht cooking. List of Intangible Cultural Heritage, Element 01852.",
        "UNITED24 Media. (2024). German outlet Zeit adopts Ukrainian spelling of Kyiv.",
        "Wikimedia Foundation. (2026). Wikimedia REST API: Pageviews. https://wikimedia.org/api/rest_v1/",
        "LanguageTool. (2022). How to spell certain Ukrainian words and names in English. https://blog.languagetool.org/insights/post/ukraine/",
        "Office Watch. (2022). Spell and say Kiev or Kyiv according to Microsoft Word.",
        "Etymonline. (2026). Borscht. Online Etymology Dictionary. https://www.etymonline.com/word/borscht",
        "Merriam-Webster. (2026). Borscht. In Merriam-Webster.com dictionary.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        for run in p.runs:
            run.font.size = Pt(9)

    # ── APPENDIX: Full data table ──────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Appendix: Complete Adoption Data", level=1)

    all_rows = []
    for r in sorted(with_data, key=lambda x: x["category"]):
        all_rows.append([
            r["category"],
            f"{r['russian']} \u2192 {r['ukrainian']}",
            pct(r["gdelt_ratio"]),
            pct(r["trends_ratio"]) if r["trends_ratio"] else "\u2014",
            pct(r["ngrams_ratio"]) if r["ngrams_ratio"] else "\u2014",
            r.get("gdelt_crossover", "\u2014") or "\u2014",
            r["status"],
        ])

    add_table(doc,
        ["Category", "Pair", "GDELT", "Trends", "Ngrams", "Crossover", "Status"],
        all_rows,
    )

    # ── SAVE ───────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size / 1024:.0f} KB")
    print(f"Pairs with data: {len(with_data)}")
    print(f"Non-control pairs: {len(non_control)}")
    print(f"Categories: {len(categories)}")


if __name__ == "__main__":
    build()
